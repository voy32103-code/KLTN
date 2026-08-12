from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPORT = Path(r"C:\Users\OS 11\Downloads\BaoCao_ReqSimulator_HUFLIT_FINAL.docx")


def replace_runs_keep_fields_and_bookmarks(paragraph, value):
    for child in list(paragraph._p):
        if child.tag.endswith("}r"):
            paragraph._p.remove(child)
    run = paragraph.add_run(value)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def set_cell(cell, value, *, header=False, status=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if status else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(value)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9.5)
    run.bold = header
    if status and value.startswith("✓"):
        run.font.color.rgb = RGBColor(0x00, 0x61, 0x00)
        run.bold = True
    if status and value.startswith("△"):
        run.font.color.rgb = RGBColor(0x9C, 0x57, 0x00)
        run.bold = True
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade(cell, color="D9EAF7"):
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), color)
    props.append(node)


def table_after(doc, caption):
    headers = ["Use case", "Module/API chính", "Artifact kiểm thử", "Trạng thái"]
    rows = [
        ("UC-01, UC-02\nĐăng ký, đăng nhập", "AuthController\nPOST /api/Auth/register; POST /api/Auth/login", "API/code đã đối chiếu; chưa có test runtime độc lập trong log hiện tại.", "△ Cần test"),
        ("UC-03, UC-04\nChọn scenario, tạo phiên", "ScenariosController; SessionsController\nGET /api/Scenarios; POST /api/Sessions", "Route, entity SimulationSession và UML đã đối chiếu; cần E2E Student flow.", "△ Cần test"),
        ("UC-05, UC-06\nXem lịch sử, gửi câu hỏi", "SessionsController; AI /api/chat\nGET messages; POST messages", "Có test retry/parsing AI service; chưa có E2E UI–API cho chat retry.", "△ Cần test"),
        ("UC-07–UC-09\nKết thúc, evaluation, report", "SessionsController; AI /api/extract, /api/evaluate", "Có finalization lease và test AI service; evaluation metric chưa có holdout run hợp lệ.", "△ Cần đánh giá"),
        ("UC-11\nReview và override", "SessionsController\nPUT /api/Sessions/review/{id}/override", "Entity LecturerOverride, sequence diagram và API route đã đối chiếu.", "✓ Đã đối chiếu"),
        ("UC-15–UC-18\nPreview, chỉnh sửa, publish", "AdminScenariosController; ScenarioVersionPublisher", "Preview–edit–publish và versioning đã đối chiếu; crawler SPA còn là giới hạn mở.", "✓ Đã đối chiếu"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate(headers):
        set_cell(table.rows[0].cells[i], value, header=True)
        shade(table.rows[0].cells[i])
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            set_cell(cells[i], value, status=(i == 3))
    widths = [2.9, 4.3, 6.6, 3.2]
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Cm(width)
    caption._p.addnext(table._tbl)


def main():
    doc = Document(REPORT)
    paragraphs = doc.paragraphs
    heading_index = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "3.4. Use case và traceability")
    body_caption = next(
        p for p in paragraphs[heading_index + 1:]
        if p.text.strip().startswith("Bảng 3.2. Ma trận truy vết")
    )
    list_entry = next(
        p for p in paragraphs[:heading_index]
        if p.text.strip().startswith("Bảng 3.2.")
    )

    # Repair the earlier list-only replacement and retain its PAGEREF field.
    for p in list(doc.paragraphs):
        if p.text.strip().startswith("Bảng 3.2 liên kết từng use case"):
            p._element.getparent().remove(p._element)
            break

    title = "Bảng 3.2. Ma trận truy vết Use Case – Module – Test Case của ReqSimulator."
    replace_runs_keep_fields_and_bookmarks(body_caption, title)
    body_caption.style = "Caption"
    replace_runs_keep_fields_and_bookmarks(list_entry, title)
    list_entry.style = "Normal"
    list_entry.paragraph_format.first_line_indent = Cm(0)

    lead = body_caption.insert_paragraph_before(
        "Bảng 3.2 làm rõ chuỗi truy vết từ use case đến module/API và artifact hiện có. "
        "Ký hiệu ✓ cho biết route, module hoặc UML đã được đối chiếu; ký hiệu △ cho biết "
        "vẫn thiếu test runtime hoặc đánh giá phù hợp. Ký hiệu này không suy ra chất lượng "
        "toàn diện của use case."
    )
    lead.style = "Normal"
    lead.paragraph_format.first_line_indent = Cm(0.75)
    lead.paragraph_format.space_after = Pt(6)
    table_after(doc, body_caption)
    doc.save(REPORT)


if __name__ == "__main__":
    main()
