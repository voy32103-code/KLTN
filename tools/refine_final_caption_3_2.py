from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt


REPORT = Path(r"C:\Users\OS 11\Downloads\BaoCao_ReqSimulator_HUFLIT_FINAL.docx")


def replace_runs_preserving_bookmarks(paragraph, value):
    # Remove runs only; list-of-tables PAGEREF bookmarks stay in the paragraph.
    for child in list(paragraph._p):
        if child.tag.endswith("}r"):
            paragraph._p.remove(child)
    run = paragraph.add_run(value)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def main():
    doc = Document(REPORT)
    caption = next(
        p for p in doc.paragraphs
        if p.text.strip().startswith("Bảng 3.2. Ma trận truy vết Use Case")
    )
    replace_runs_preserving_bookmarks(
        caption,
        "Bảng 3.2. Ma trận truy vết Use Case – Module – Test Case của ReqSimulator.",
    )
    caption.style = "Caption"

    # The explanatory sentence belongs immediately before the caption, not in
    # the table-of-tables entry.
    previous = caption.insert_paragraph_before(
        "Bảng 3.2 liên kết từng use case với controller, service và test case tương ứng. "
        "Use case chưa có test runtime độc lập được ghi nhận là cần đánh giá thêm; "
        "chúng không được xem là đã được chứng minh bằng kiểm thử kỹ thuật."
    )
    previous.style = "Normal"
    previous.paragraph_format.first_line_indent = Cm(0.75)
    previous.paragraph_format.space_after = Pt(6)
    doc.save(REPORT)


if __name__ == "__main__":
    main()
