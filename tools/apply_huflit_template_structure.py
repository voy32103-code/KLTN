"""Apply the usable structural conventions of BaoCao_KLTN.docx to ReqSimulator.

The supplied file is a Word-format template but contains a different thesis.  This
script keeps ReqSimulator content and adds the template's front-matter convention:
table/figure lists, bookmarks with PAGEREF fields, and styles suitable for the
later Word section/page-number pass.
"""

from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\OS 11\Downloads\BaoCao_ReqSimulator_HUFLIT_IEEE_rebuild_v1.docx")
TARGET = Path(r"C:\Users\OS 11\Downloads\BaoCao_ReqSimulator_HUFLIT_FINAL.docx")
CAPTION_PATTERN = re.compile(r"^(Bảng|Hình)\s+([A-Z]?\.?\d+(?:\.\d+)*)\.\s+(.+)$")


def make_paragraph(parent, value, style):
    element = OxmlElement("w:p")
    paragraph = Paragraph(element, parent)
    paragraph.style = style
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(value)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    return paragraph


def add_bookmark(paragraph, name, bookmark_id):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_page_ref(paragraph, bookmark):
    ppr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "8900")
    tabs.append(tab)
    ppr.append(tabs)
    paragraph.add_run("\t")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), f"PAGEREF {bookmark} \\h")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "0"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def block_before(anchor, heading, items):
    parent = anchor._parent
    heading_p = make_paragraph(parent, heading, "Heading 1")
    heading_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    anchor._p.addprevious(heading_p._p)
    for caption, bookmark in items:
        entry = make_paragraph(parent, caption, "Normal")
        add_page_ref(entry, bookmark)
        anchor._p.addprevious(entry._p)


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    doc = Document(SOURCE)

    tables = []
    figures = []
    bookmark_id = 100
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        match = CAPTION_PATTERN.match(text)
        if not match:
            continue
        kind = match.group(1)
        # The explanatory lead-in before Bảng 1.1 is not a caption because it
        # does not match this pattern.  Captions are explicitly styled here.
        paragraph.style = "Caption"
        name = ("tbl_" if kind == "Bảng" else "fig_") + str(bookmark_id)
        add_bookmark(paragraph, name, bookmark_id)
        bookmark_id += 1
        item = (text, name)
        (tables if kind == "Bảng" else figures).append(item)

    anchor = next(
        p for p in doc.paragraphs
        if p.text.strip() == "DANH MỤC CÁC KÝ HIỆU, CHỮ VIẾT TẮT VÀ THUẬT NGỮ"
    )
    # Add tables first, then figures, preserving the order expected by the template.
    block_before(anchor, "DANH MỤC CÁC BẢNG", tables)
    block_before(anchor, "DANH MỤC CÁC HÌNH VẼ, ĐỒ THỊ", figures)

    # Match the template's body typography and heading rhythm without copying
    # its unrelated content.
    for name, size, before, after, line in (
        ("Heading 1", 14, 12, 6, 1.5),
        ("Heading 2", 13, 10, 6, 1.5),
        ("Heading 3", 13, 8, 4, 1.5),
    ):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = name != "Heading 3"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = line
        style.paragraph_format.keep_with_next = True

    doc.save(TARGET)
    print(TARGET)
    print(f"tables={len(tables)} figures={len(figures)}")


if __name__ == "__main__":
    main()
