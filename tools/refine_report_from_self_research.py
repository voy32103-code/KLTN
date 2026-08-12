"""Replace generic thesis tables with statements traceable to project artifacts.

Input:  the report that contains the first research-rationale revision.
Output: a separate report so the previous Word version remains recoverable.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\OS 11\Downloads\BaoCao_ReqSimulator_HUFLIT_IEEE_final_research_rationale_updated.docx")
TARGET = Path(r"C:\Users\OS 11\Downloads\BaoCao_ReqSimulator_HUFLIT_IEEE_final_self_research_evidence_updated.docx")


def set_paragraph(paragraph, value, *, italic=False, centered=False):
    paragraph._p.clear_content()
    run = paragraph.add_run(value)
    run.font.size = Pt(11)
    run.italic = italic
    if centered:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_cell(cell, value, *, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(value)
    run.font.size = Pt(8.5)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), color)
    props.append(node)


def replace_table(table, headers, rows):
    while len(table.rows) < len(rows) + 1:
        table.add_row()
    while len(table.rows) > len(rows) + 1:
        table._tbl.remove(table.rows[-1]._tr)

    for col, header in enumerate(headers):
        set_cell(table.rows[0].cells[col], header, bold=True)
        shade(table.rows[0].cells[col], "D9EAF7")
    for row_index, values in enumerate(rows, start=1):
        for col, value in enumerate(values):
            set_cell(table.rows[row_index].cells[col], value)


def paragraph_after_table(table, value):
    node = OxmlElement("w:p")
    table._tbl.addnext(node)
    paragraph = Paragraph(node, table._parent)
    set_paragraph(paragraph, value)
    return paragraph


def find_paragraph(doc, begins):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(begins):
            return paragraph
    raise RuntimeError(f"Không tìm thấy đoạn: {begins}")


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    doc = Document(SOURCE)

    # Table 1.1 is the first newly added table, after the abbreviation table.
    intro = find_paragraph(doc, "Bảng 1.1 là đối chiếu chức năng")
    set_paragraph(
        intro,
        "Bảng 1.1 là đối chiếu phạm vi và cơ chế do người thực hiện tổng hợp từ "
        "Báo cáo tự nghiên cứu ReqSimulator và artifact thiết kế. Bảng không phải "
        "benchmark thương mại, cũng không đánh giá chất lượng của sản phẩm bên ngoài.",
    )
    caption = find_paragraph(doc, "Bảng 1.1. Đối chiếu chức năng")
    set_paragraph(
        caption,
        "Bảng 1.1. Phạm vi khác biệt của ReqSimulator so với các cách tiếp cận liên quan",
        italic=True,
        centered=True,
    )
    replace_table(
        doc.tables[2],
        ["Cách tiếp cận", "Cách xử lý tri thức / hội thoại", "Phù hợp và giới hạn với bài toán", "Quyết định trong ReqSimulator"],
        [
            (
                "Chatbot LLM tự do",
                "Sinh phản hồi từ prompt và lịch sử hội thoại mở.",
                "Hữu ích để giải thích hoặc trao đổi tự do, nhưng không có ground truth ẩn theo scenario để truy vết người học đã khai thác gì.",
                "Gắn persona với scenario, transcript và hidden requirement có version; hệ thống không coi câu trả lời LLM là chuẩn chấm điểm.",
            ),
            (
                "RAG truyền thống",
                "Truy hồi các đoạn tài liệu liên quan rồi cấp context rộng cho LLM trả lời câu hỏi.",
                "Phù hợp hỏi–đáp trên kho tri thức; không bảo đảm thông tin chỉ được tiết lộ sau đúng loại câu hỏi trong một bài luyện elicitation.",
                "Dùng deterministic gating: chỉ requirement thỏa question type, topic và gate mới được đưa vào phạm vi diễn đạt của LLM.",
            ),
            (
                "Công cụ quản lý requirement doanh nghiệp",
                "Lưu, phân công, theo dõi thay đổi và liên kết requirement sau khi nội dung đã được ghi nhận.",
                "Không nhắm đến việc mô phỏng stakeholder để sinh viên luyện đặt câu hỏi trước giai đoạn đặc tả.",
                "ReqSimulator chỉ hỗ trợ giai đoạn luyện elicitation; không tuyên bố thay thế DOORS, Jira hoặc công cụ quản lý dự án.",
            ),
            (
                "Quiz / biểu mẫu tĩnh",
                "Chấm câu trả lời theo câu hỏi cố định và ít phụ thuộc trạng thái trước đó.",
                "Khó phản ánh hội thoại nhiều lượt, câu trả lời mơ hồ và việc mở thông tin có điều kiện.",
                "Lưu trạng thái persona, metadata câu hỏi và hội thoại nhiều lượt; sau phiên mới chạy extraction và matching.",
            ),
            (
                "ReqSimulator",
                "Scenario phiên bản hóa + persona + controlled disclosure + extraction AAOC + matching với ground truth ẩn.",
                "Tạo môi trường lặp lại để luyện một phiên phỏng vấn và để giảng viên review/override.",
                "Phạm vi là prototype đào tạo và kiểm chứng kỹ thuật; hiệu quả giáo dục, Precision/Recall/F1 và kết quả A/B chỉ được kết luận khi có evaluation hợp lệ.",
            ),
        ],
    )

    tech_intro = find_paragraph(doc, "Các công nghệ được chọn theo ràng buộc")
    set_paragraph(
        tech_intro,
        "Lựa chọn công nghệ phản ánh trade-off thực tế trong quá trình tự triển khai: "
        "tách public API, AI orchestration và dữ liệu nghiệp vụ; giảm dependency không cần thiết; "
        "đồng thời giữ khả năng kiểm thử và thay provider AI. Vì vậy, bảng nêu cả lý do chọn "
        "lẫn đánh đổi, không xem bất kỳ công nghệ nào là mặc định tốt hơn mọi lựa chọn khác.",
    )
    tech_caption = find_paragraph(doc, "Bảng 4.1. Lý do lựa chọn công nghệ")
    set_paragraph(tech_caption, "Bảng 4.1. Quyết định công nghệ và đánh đổi trong quá trình tự triển khai", italic=True, centered=True)
    replace_table(
        doc.tables[7],
        ["Thành phần", "Lựa chọn", "Lý do chọn trong ReqSimulator", "Đánh đổi / lý do không mở rộng"],
        [
            (
                "Frontend",
                "Vite + Vanilla TypeScript + Chart.js",
                "MVP cần build nhanh, dependency ít, payload API có kiểm tra kiểu và dashboard có biểu đồ cơ bản.",
                "Khi UI lớn hơn, code cần tiếp tục mô-đun hóa và bổ sung test; không suy ra kiến trúc này phù hợp mọi SPA.",
            ),
            (
                "Backend nghiệp vụ",
                "ASP.NET Core 9 + EF Core/Npgsql",
                "Phù hợp code C# hiện có cho JWT/RBAC, ownership, transaction và quan hệ User–Session–Evaluation.",
                "Tạo boundary C#–Python nên DTO/contract test phải được duy trì thay vì gộp toàn bộ logic vào một service.",
            ),
            (
                "AI service",
                "FastAPI + Pydantic",
                "Tách gọi GenAI, gating, extraction, retry và kiểm tra structured output khỏi public API; schema hỗ trợ từ chối payload lỗi.",
                "Hai hệ type TypeScript/C#/Python làm tăng chi phí đồng bộ contract và test liên dịch vụ.",
            ),
            (
                "Persistence",
                "PostgreSQL + JSONB",
                "Dữ liệu phiên có quan hệ rõ ràng, còn scenario state/structured output cần trường linh hoạt để lưu audit.",
                "Schema và migration phải được kiểm soát; JSONB không thay thế ràng buộc cho các trường nghiệp vụ cốt lõi.",
            ),
            (
                "Embedding / LLM",
                "Gemini API + NumPy; adapter đa provider",
                "Embedding cục bộ dùng PyTorch/SentenceTransformers từng vượt giới hạn image/RAM của môi trường triển khai nhỏ; API giảm gánh nặng tài nguyên.",
                "Đổi lại phụ thuộc mạng, quota và model version; metadata model/threshold cần lưu đầy đủ trước evaluation chính thức.",
            ),
        ],
    )

    issue_intro = find_paragraph(doc, "Các vấn đề dưới đây xuất hiện")
    set_paragraph(
        issue_intro,
        "Bảng 4.2 ghi các sự cố và ràng buộc xuất hiện trong artifact triển khai. Mỗi dòng "
        "phân biệt triệu chứng, nguyên nhân, phần đã sửa và việc còn phải kiểm chứng; không dùng "
        "roadmap để tuyên bố chức năng đã hoàn tất.",
    )
    issue_caption = find_paragraph(doc, "Bảng 4.2. Vấn đề triển khai")
    set_paragraph(issue_caption, "Bảng 4.2. Vấn đề thực tế, nguyên nhân và biện pháp xử lý", italic=True, centered=True)
    replace_table(
        doc.tables[8],
        ["Vấn đề thực tế", "Nguyên nhân gốc", "Biện pháp đã áp dụng / artifact", "Việc còn phải kiểm chứng"],
        [
            (
                "Cập nhật scenario có thể vướng khóa ngoại của kết quả cũ",
                "Đồng bộ bằng cách xóa/tạo lại hidden requirement trong khi RequirementMatch lịch sử vẫn tham chiếu dữ liệu cũ.",
                "ScenarioVersionPublisher tạo version, persona và hidden requirement mới; version cũ chỉ superseded. Publish dùng transaction/advisory lock.",
                "Cần integration test PostgreSQL thật cho publisher đồng thời và migration dữ liệu cũ.",
            ),
            (
                "Kết thúc phiên đồng thời tạo evaluation trùng hoặc score fallback giả",
                "External AI call dài nhưng không có quyền sở hữu finalization rõ ràng; retry/double-submit cạnh tranh nhau.",
                "Thêm finalization lease + unique evaluation; AI fallback trả 503 và không persist như kết quả chính thức.",
                "Cần test lease expiry, crash và retry trong môi trường tích hợp.",
            ),
            (
                "Crawler tạo scenario khó dùng khi publish trực tiếp output AI",
                "Ingestion và publish từng là một bước, nên admin không có nơi sửa output xác suất trước khi ghi DB.",
                "Tách preview–edit–publish để admin kiểm tra context, hidden requirement và gate trước publish.",
                "Hai đường direct-publish cũ cần được deprecate sau khi xác nhận không còn consumer.",
            ),
            (
                "Embedding cục bộ vượt tài nguyên môi trường nhỏ",
                "PyTorch/SentenceTransformers làm image và RAM lớn so với quota triển khai thử nghiệm.",
                "Chuyển sang Gemini Embedding API và giữ NumPy cho phép tính vector.",
                "Cần provenance metadata, quota handling và evaluation độc lập trước khi kết luận matching tốt hơn.",
            ),
            (
                "20 test async không chạy được",
                "Môi trường dev thiếu pytest-asyncio nên pytest không thực thi coroutine được đánh dấu pytest.mark.asyncio.",
                "Bổ sung requirements-dev.txt, cài pytest-asyncio==1.4.0 và chạy lại: 92 passed, 11 warnings; log 09/08/2026.",
                "Dependency và log cần được commit/tag cùng source để tái lập snapshot nộp báo cáo.",
            ),
            (
                "Không có cơ sở hợp lệ để kết luận A/B feedback",
                "Chưa có consent được phê duyệt; schema chưa lưu consent version/thời điểm, cỡ mẫu và phân tích chưa khóa trước.",
                "Lập readiness gate và không chạy/không công bố variant thắng trên dữ liệu người dùng thật.",
                "Chỉ triển khai khi hoàn tất consent, schema consent, sampling plan và phân tích định trước.",
            ),
        ],
    )
    paragraph_after_table(
        doc.tables[8],
        "Các giới hạn còn mở được giữ nguyên: crawler HTTP chưa render đầy đủ SPA; gửi tin nhắn chưa idempotent hoàn toàn; "
        "metadata evaluation chưa lưu đủ model/prompt/threshold; stateUpdate cần tiếp tục được loại khỏi public contract; video chưa upload end-to-end; "
        "test frontend và integration database còn mỏng. Đây là backlog kỹ thuật, không được mô tả là chức năng đã hoàn tất.",
    )

    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
