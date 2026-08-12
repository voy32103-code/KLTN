from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


SOURCE = Path(r"C:\Users\OS 11\Downloads\BaoCao_ReqSimulator_HUFLIT_IEEE_final_with_placeholder.docx")
TARGET = Path(r"C:\Users\OS 11\Downloads\BaoCao_ReqSimulator_HUFLIT_IEEE_final_with_placeholder_validation_updated.docx")


def find_paragraph(doc, prefix):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Missing paragraph: {prefix}")


def replace_paragraph(paragraph, text, italic=False, centered=False):
    paragraph._p.clear_content()
    run = paragraph.add_run(text)
    run.italic = italic
    run.font.size = Pt(11)
    if centered:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def table_text(table):
    return ' '.join(cell.text for row in table.rows for cell in row.cells)


def set_cell(cell, text):
    cell.text = ''
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(9)


def main():
    doc = Document(SOURCE)
    table_51 = next(t for t in doc.tables if 'AI service unit tests' in table_text(t))
    for row in table_51.rows[1:]:
        topic = row.cells[0].text.strip()
        if topic == 'Snapshot mã nguồn':
            set_cell(row.cells[1], 'Commit cơ sở 7efad62 (08/08/2026); dependency test bổ sung tại ai-service/requirements-dev.txt; log: docs/audits_evaluation/verification_2026-08-09.md')
            set_cell(row.cells[2], 'Mốc đối chiếu source; dependency test phải được commit/tag trước bản nộp chính thức.')
        elif topic == 'AI service unit tests':
            set_cell(row.cells[1], 'cd ai-service; .\\.venv\\Scripts\\python.exe -m pytest -q')
            set_cell(row.cells[2], '92 passed, 11 warnings, 1.23 s sau khi cài pytest-asyncio==1.4.0. Warning deprecation từ thư viện; không có failure.')

    replace_paragraph(
        find_paragraph(doc, 'Các kiểm chứng kỹ thuật được đóng gói'),
        'Các kiểm chứng kỹ thuật được đóng gói thành artifact để người đọc có thể tái chạy và đối chiếu. Bảng 5.1 ghi kết quả tại snapshot đã nêu; sau khi bổ sung runtime test async, AI service đạt 92 test passed. Log đầy đủ lưu tại docs/audits_evaluation/verification_2026-08-09.md.'
    )
    replace_paragraph(
        find_paragraph(doc, 'Bảng 5.1.'),
        'Bảng 5.1. Artifact kiểm chứng build/test tại snapshot 7efad62 (cập nhật chạy ngày 09/08/2026)',
        italic=True, centered=True,
    )
    replace_paragraph(
        find_paragraph(doc, 'Bộ dữ liệu hiện có là corpus pilot'),
        'Dataset pilot-v1 đã được khóa bằng manifest docs/research_academic/pilot_dataset/dataset_lock_pilot_v1.json: 10 transcript tổng hợp, checksum từng transcript/annotation, scenario checksum, split 8 session calibration và 2 session holdout không trùng transcript. Annotation hiện là version 1, chưa có hai annotator độc lập và adjudication. Lần chạy raw LLM extraction trên holdout chỉ được thực hiện khi được phép gửi transcript tổng hợp tới provider; vì vậy chưa điền Precision, Recall, F1 hoặc confusion matrix vào kết quả cuối cùng.'
    )
    replace_paragraph(
        find_paragraph(doc, 'Bảng 5.2.'),
        'Bảng 5.2. Ma trận quy trình đánh giá extraction/matching (pilot-v1 đã khóa; kết quả cuối chờ dual review và raw LLM run)',
        italic=True, centered=True,
    )
    replace_paragraph(
        find_paragraph(doc, 'Hệ thống đã có cơ chế lưu feedback survey'),
        'Hệ thống đã có cơ chế gán variant theo session, lưu survey và tổng hợp Helpfulness, Actionability, No-answer-leak. Tuy nhiên, artifact kiểm tra ngày 09/08/2026 xác nhận chưa có consent được phê duyệt và schema chưa lưu trạng thái/phiên bản consent. Vì vậy, A/B chưa được phép chạy với người dùng thật; Bảng 5.3 chỉ là kế hoạch đánh giá và không kết luận variant nào hữu ích hơn.'
    )

    appendix_a = find_paragraph(doc, 'Phụ lục A liệt kê artifact')
    replace_paragraph(
        appendix_a,
        'Phụ lục A liệt kê artifact kiểm chứng dùng cho Bảng 5.1 và các gate thực nghiệm. Artifact lưu trong repository để người chấm đối chiếu snapshot, lệnh chạy, dataset lock và điều kiện đạo đức của A/B.'
    )
    manifest_table = next(t for t in doc.tables if 'ART-01' in table_text(t) and 'ART-04' in table_text(t))
    new_row = manifest_table.add_row().cells
    set_cell(new_row[0], 'ART-05')
    set_cell(new_row[1], 'Dataset lock pilot-v1 và A/B readiness gate.')
    set_cell(new_row[2], 'pilot_dataset/dataset_lock_pilot_v1.json; audits_evaluation/ab_experiment_readiness_2026-08-09.md — chưa đủ điều kiện kết luận thực nghiệm.')

    settings = doc.settings.element
    doc.save(TARGET)
    print(TARGET)


if __name__ == '__main__':
    main()
