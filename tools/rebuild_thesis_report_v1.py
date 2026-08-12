"""First clean rebuild of the ReqSimulator thesis.

The script intentionally keeps the original cover, UML images, and verified artifacts.
It rewrites generic prose, strengthens the research thread, and applies one Word style
system. It does not fabricate evaluation results.
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\OS 11\Downloads\BaoCao_ReqSimulator_HUFLIT_IEEE_final_self_research_evidence_updated.docx")
TARGET = Path(r"C:\Users\OS 11\Downloads\BaoCao_ReqSimulator_HUFLIT_IEEE_rebuild_v1.docx")


def find(doc, prefix, occurrence=0):
    matches = [p for p in doc.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) <= occurrence:
        raise RuntimeError(f"Không tìm thấy đoạn {prefix!r} ({occurrence})")
    return matches[occurrence]


def find_body(doc, prefix):
    """Find a body paragraph, excluding the result text of the Word TOC field."""
    for paragraph in doc.paragraphs:
        if paragraph.style.name.lower().startswith("toc"):
            continue
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Không tìm thấy body paragraph {prefix!r}")


def replace(paragraph, value, style=None, *, italic=False, centered=False):
    paragraph._p.clear_content()
    if style:
        paragraph.style = style
    run = paragraph.add_run(value)
    run.italic = italic
    if centered:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return paragraph


def add_after(anchor, value, style=None, *, italic=False, centered=False):
    element = OxmlElement("w:p")
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    replace(paragraph, value, style, italic=italic, centered=centered)
    return paragraph


def set_cell(cell, value, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(value)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade(cell, color="D9EAF7"):
    props = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), color)
    props.append(element)


def configure_table(table, headers=None):
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
        if row_index == 0:
            row._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    if headers:
        for i, value in enumerate(headers):
            set_cell(table.rows[0].cells[i], value, bold=True)
            shade(table.rows[0].cells[i])


def insert_table_after(doc, anchor, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    configure_table(table, headers)
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            set_cell(cells[i], value)
    anchor._p.addnext(table._tbl)
    return table


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(0.75)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size in (("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 13)):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.line_spacing = 1.15
    # Chapter starts are controlled paragraph-by-paragraph below. Applying a page
    # break to the style also breaks front matter such as "LỜI CẢM ƠN".
    doc.styles["Heading 1"].paragraph_format.page_break_before = False

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(11)
    caption.font.italic = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.first_line_indent = Cm(0)

    # The generated TOC inherits Normal in the source file. Give it an explicit
    # compact style so the last entry is not stranded on a third page.
    for name, size in (("toc 1", 11), ("toc 2", 11), ("toc 3", 10)):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.first_line_indent = Cm(0)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        # The existing HUFLIT cover was laid out for a 3 cm left margin. Keep
        # it on one page until an official faculty template is supplied.
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)


def improve_front_matter(doc):
    replace(
        find(doc, "Lời đầu tiên,"),
        "Em trân trọng cảm ơn ThS. Đặng Thị Kim Giao đã hướng dẫn, góp ý về phạm vi đề tài "
        "và yêu cầu đối với báo cáo khóa luận. Những phản hồi của Cô giúp em xác định rõ "
        "phần nào cần có artifact kiểm chứng và phần nào chỉ được trình bày là kế hoạch đánh giá.",
    )
    replace(
        find(doc, "Trong suốt thời gian qua,"),
        "Em cảm ơn quý thầy cô Khoa Công nghệ Thông tin, Trường Đại học Ngoại ngữ - Tin học "
        "Thành phố Hồ Chí Minh đã cung cấp nền tảng kiến thức về phân tích, thiết kế và phát triển phần mềm. "
        "Em cũng cảm ơn gia đình và bạn bè đã hỗ trợ trong thời gian thực hiện đề tài.",
    )
    replace(
        find(doc, "Dù đã có nhiều cố gắng,"),
        "Báo cáo còn những giới hạn đã được nêu rõ ở Chương 5 và Chương 6. Em xin tiếp thu ý kiến "
        "của Hội đồng để tiếp tục hoàn thiện hệ thống và phương pháp đánh giá.",
    )
    # Remove the two redundant acknowledgement paragraphs while retaining page structure.
    for prefix in ("Em cũng xin chân thành cảm ơn", "Cuối cùng, em xin cảm ơn"):
        paragraph = find(doc, prefix)
        paragraph._element.getparent().remove(paragraph._element)

    replace(
        find(doc, "Khóa luận trình bày quá trình"),
        "Khóa luận trình bày việc phân tích, thiết kế và hiện thực ReqSimulator, một hệ thống mô phỏng "
        "phỏng vấn stakeholder để sinh viên luyện kỹ năng khai thác yêu cầu phần mềm. Người học làm việc "
        "với scenario có version, persona stakeholder và danh sách yêu cầu ẩn; hệ thống lưu transcript, "
        "trích xuất requirement có cấu trúc và tạo báo cáo sau phiên.",
    )
    replace(
        find(doc, "Đề tài xây dựng kiến trúc gồm"),
        "Kiến trúc hiện thực gồm web client Vite/TypeScript, API nghiệp vụ ASP.NET Core, AI service FastAPI "
        "và PostgreSQL. AI service thực hiện controlled disclosure, extraction theo Actor–Action–Object–Condition "
        "(AAOC), matching với ground truth và tạo feedback không đưa nguyên văn đáp án scenario xuống client. "
        "Giảng viên có luồng review và override để không biến kết quả do AI sinh thành điểm chính thức.",
    )
    replace(
        find(doc, "Phần đánh giá được xây dựng"),
        "Bằng chứng của khóa luận được giới hạn ở source code, UML, test và log chạy kỹ thuật. Dataset pilot-v1 "
        "đã có manifest và split theo session, nhưng annotation chưa dual review/adjudication; raw LLM holdout run "
        "chưa được phê duyệt. Vì vậy báo cáo không công bố Precision, Recall, F1, hiệu quả học tập hoặc kết quả A/B. "
        "Các nội dung đó được trình bày như giao thức đánh giá cần hoàn tất sau.",
    )
    replace(find(doc, "Từ khóa:"), "Từ khóa: khai thác yêu cầu; phỏng vấn stakeholder; mô phỏng theo kịch bản; controlled disclosure; AAOC.")
    replace(
        find(doc, "This thesis presents"),
        "This thesis presents ReqSimulator, an interactive system for practising software requirements elicitation "
        "through simulated stakeholder interviews. A student works with versioned scenarios, stakeholder personas, "
        "and hidden requirements. The system stores the transcript, extracts structured requirements, matches them "
        "against scenario ground truth, and produces post-session feedback without exposing the answer key.",
    )
    replace(
        find(doc, "Keywords:"),
        "Keywords: requirements elicitation; stakeholder interview; scenario-based simulation; controlled disclosure; AAOC.",
    )


def strengthen_chapter_two(doc):
    replace(
        find(doc, "Khai thác yêu cầu là hoạt động"),
        "Khai thác yêu cầu là quá trình tìm, làm rõ và thống nhất nhu cầu của các bên liên quan để tạo đầu vào "
        "có thể kiểm tra cho đặc tả. ISO/IEC/IEEE 29148 yêu cầu yêu cầu phải rõ nguồn gốc, nhất quán và truy vết được; "
        "các tổng quan về kỹ nghệ yêu cầu xem phỏng vấn, quan sát và phân tích tài liệu là các kỹ thuật trọng tâm [1], [2]. "
        "Trong ReqSimulator, đơn vị luyện tập là một phiên phỏng vấn. Vì vậy trọng tâm không phải là nhớ định nghĩa FR/NFR/BR, "
        "mà là nhận ra thông tin còn thiếu và đặt câu hỏi làm rõ có căn cứ.",
    )
    p = find(doc, "Một buổi phỏng vấn có chất lượng")
    replace(
        p,
        "Trong phỏng vấn, câu trả lời đầu tiên thường chưa đủ để xác định actor, thao tác, đối tượng, điều kiện và quy tắc nghiệp vụ. "
        "Sommerville phân biệt yêu cầu người dùng với yêu cầu hệ thống; Zowghi và Coulin nhấn mạnh ảnh hưởng của giao tiếp, "
        "mâu thuẫn mục tiêu và tri thức ngầm của stakeholder [3], [4]. Thiết kế ReqSimulator vì thế lưu hội thoại theo phiên "
        "và đánh giá phát biểu rút ra từ transcript, thay vì coi số lượt hỏi là chỉ dấu chất lượng.",
    )
    add_after(
        p,
        "Từ góc nhìn thiết kế, một câu hỏi chỉ có giá trị khi làm giảm một điểm chưa rõ trong scenario. Hệ thống phân loại "
        "câu hỏi theo mức vague, on-topic, specific và conditional; kết quả phân loại không thay thế phán đoán của giảng viên, "
        "nhưng là điều kiện để kiểm soát mức thông tin mà stakeholder ảo có thể tiết lộ.",
    )

    replace(
        find(doc, "Kịch bản tạo một bối cảnh"),
        "Kịch bản xác định miền nghiệp vụ, vai trò stakeholder, mục tiêu, ràng buộc và tập yêu cầu cần được khai thác. "
        "Scenario-based design cho phép người học lặp lại một tình huống trong điều kiện tương đối ổn định, quan sát hệ quả của lựa chọn "
        "và so sánh các lần thực hiện [5]. Trong hệ thống, scenario được version hóa để một session đã bắt đầu không bị chấm bằng cấu hình "
        "mới do quản trị viên publish sau đó.",
    )
    p = find(doc, "Phản hồi chỉ có ích")
    replace(
        p,
        "Phản hồi cần nối mục tiêu, trạng thái hiện tại và hành động tiếp theo. Hattie và Timperley mô tả ba câu hỏi: mục tiêu là gì, "
        "người học đang ở đâu và bước tiếp theo nên làm gì [6]. Với ReqSimulator, feedback chỉ nhận category, trạng thái match và coverage; "
        "đầu ra phải hướng người học đến vùng cần làm rõ chứ không phát lại nguyên văn hidden requirement.",
    )
    add_after(
        p,
        "Do chưa có user study hợp lệ, việc đưa feedback vào hệ thống là một quyết định thiết kế dựa trên nguyên tắc phản hồi; "
        "khóa luận không suy ra rằng feedback đã cải thiện kỹ năng của người học.",
    )

    replace(
        find(doc, "Mô hình ngôn ngữ lớn có thể"),
        "Mô hình ngôn ngữ lớn có thể duy trì hội thoại theo vai và chuyển văn bản tự do thành dữ liệu có cấu trúc. Tuy nhiên, đầu ra "
        "có tính xác suất và có thể tạo thông tin không được hỗ trợ bởi context [7], [8]. Nếu đưa toàn bộ ground truth vào prompt hoặc "
        "trả raw internal state về client, hệ thống có nguy cơ làm lộ đáp án của bài luyện.",
    )
    p = find(doc, "Ở mức hiện thực, prompt")
    replace(
        p,
        "ReqSimulator không dùng RAG mở như một chatbot hỏi–đáp kho tri thức. Deterministic gate lựa chọn requirement nào có thể được "
        "dùng để diễn đạt dựa trên question type, topic và điều kiện disclosure; LLM chỉ diễn đạt trong phạm vi đó. Structured output "
        "được kiểm tra bằng schema, còn giảng viên vẫn có quyền review/override. Cơ chế này làm rõ ranh giới giữa nội dung do scenario kiểm soát "
        "và câu trả lời do mô hình tạo.",
    )

    replace(
        find(doc, "Khả năng truy vết giúp"),
        "Khả năng truy vết liên kết một yêu cầu với nguồn phát biểu, version scenario, lần xử lý và kết quả đánh giá. Chuẩn 29148 yêu cầu "
        "duy trì liên kết giữa nhu cầu, yêu cầu và artifact kỹ thuật; biểu diễn câu bằng vector có thể hỗ trợ tìm các phát biểu gần nghĩa [1], [9]. "
        "Trong ReqSimulator, extracted requirement được liên kết với transcript, session, scenario version và hidden requirement tương ứng.",
    )
    replace(
        find(doc, "Hệ thống chuẩn hóa phát biểu"),
        "Để giảm việc matching chỉ dựa vào từ giống nhau, hệ thống chuẩn hóa phát biểu theo Actor–Action–Object–Condition (AAOC), đồng thời lưu "
        "dữ liệu thô và dữ liệu đã chuẩn hóa. Cách tách này cho phép người review xem lại vì sao một match được đề xuất, nhưng không tự chứng minh "
        "rằng phép chuẩn hóa đúng trên mọi scenario hoặc ngôn ngữ.",
    )

    replace(
        find(doc, "Matching ánh xạ"),
        "Matching ánh xạ requirement người học diễn đạt với hidden requirement của scenario. Embedding tạo ứng viên gần nghĩa, nhưng không xác lập "
        "đúng/sai cho miền nghiệp vụ [9]. Do đó, pipeline trước hết lọc theo type, action và object; sau đó dùng actor, action, object và condition "
        "để tính rubric AAOC. Gán một–một nhằm tránh việc một phát biểu quá rộng được dùng để tính coverage cho nhiều ground truth.",
    )
    p = find(doc, "ReqSimulator phân loại kết quả")
    replace(
        p,
        "Kết quả matching được phân loại exact, semantic, partial hoặc missed. Exact yêu cầu các thành phần cốt lõi tương ứng; semantic cho phép khác "
        "cách diễn đạt nhưng không đổi nội dung; partial ghi nhận thông tin có ích nhưng thiếu thành phần quan trọng; missed là không có bằng chứng đủ mạnh. "
        "Threshold và rubric chỉ có thể được chốt sau dual annotation, adjudication và đo độ đồng thuận, ví dụ bằng kappa của Cohen [10].",
    )
    add_after(
        p,
        "Trong phiên bản báo cáo này, các trọng số AAOC là tham số thiết kế để kiểm thử luồng matching; chúng chưa phải tham số thực nghiệm đã hiệu chỉnh.",
    )

    replace(
        find(doc, "ReqSimulator tách giao diện web"),
        "ReqSimulator tách giao diện web, API nghiệp vụ, AI service và cơ sở dữ liệu để tách trách nhiệm: client hiển thị luồng; backend xác thực, kiểm tra quyền "
        "và lưu dữ liệu; AI service xử lý prompt, extraction và evaluation. JWT chỉ là khuôn dạng truyền claims; chữ ký, thời hạn và ngữ cảnh sử dụng phải được kiểm tra "
        "trước khi cấp quyền [11].",
    )
    p = find(doc, "Các yêu cầu an toàn trong hệ thống")
    replace(
        p,
        "Các kiểm soát liên quan gồm RBAC, ownership check, khóa nội bộ giữa backend và AI service, không gửi hidden requirement xuống client và lưu thông tin cần audit. "
        "Những kiểm soát này phù hợp với nhóm xác thực, phân quyền và logging của OWASP ASVS [12], nhưng chưa thay thế penetration test hoặc production threat model.",
    )
    table_caption = add_after(p, "Bảng 2.1. Nguyên tắc thiết kế dùng để giới hạn rủi ro của AI trong hệ thống", "Caption")
    insert_table_after(
        doc,
        table_caption,
        ["Rủi ro", "Quyết định thiết kế", "Giới hạn của quyết định"],
        [
            ("LLM tự bịa hoặc trả ngoài context", "Gate xác định dữ kiện được phép; schema/retry kiểm tra structured output.", "Không thay thế evaluation độ chính xác trên dataset đã gán nhãn."),
            ("Lộ đáp án scenario", "Không trả hidden requirement/newly revealed state; feedback dùng category và trạng thái match.", "Cần kiểm thử contract và user study no-answer-leak."),
            ("Điểm không tái lập", "Version scenario, manifest dataset, rubric và log test được lưu như artifact.", "Metadata model/prompt/threshold chưa hoàn chỉnh ở mọi evaluation cũ."),
            ("AI ghi đè phán đoán giảng viên", "Lecturer/Admin có luồng review và override.", "Chưa đo độ đồng thuận giữa reviewer và hệ thống."),
        ],
    )


def improve_chapter_five(doc):
    replace(find_body(doc, "5.1. Mục tiêu kiểm thử"), "5.1. Kiểm thử kỹ thuật và phạm vi bằng chứng", "Heading 2")
    replace(
        find(doc, "Các kiểm chứng kỹ thuật được đóng gói"),
        "Mục này chỉ báo cáo kiểm thử kỹ thuật có thể truy về artifact. Test xanh cho biết các case có trong suite đã chạy tại snapshot đã nêu; "
        "không phải bằng chứng rằng toàn bộ hệ thống không còn lỗi, không phải benchmark hiệu năng và cũng không phải bằng chứng hiệu quả học tập. "
        "Bảng 5.1 ghi commit cơ sở, dependency, lệnh chạy và log để người đọc có thể đối chiếu.",
    )
    replace(find(doc, "Bảng 5.1. Artifact kiểm chứng"), "Bảng 5.1. Artifact kiểm chứng kỹ thuật tại snapshot đã nêu", "Caption")
    replace(find_body(doc, "5.2. Đánh giá pipeline"), "5.2. Giao thức đánh giá extraction và matching", "Heading 2")
    replace(
        find(doc, "Dataset pilot-v1 đã được khóa"),
        "Pilot-v1 đã được khóa bằng manifest `dataset_lock_pilot_v1.json`: 10 transcript tổng hợp, checksum transcript/annotation, checksum scenario "
        "và split 8 session calibration – 2 session holdout không trùng transcript hoặc người học. Tuy nhiên annotation mới ở version 1; chưa có hai annotator "
        "độc lập và adjudication. Raw LLM extraction trên holdout chỉ được thực hiện khi có quyền gửi transcript tổng hợp tới provider. Vì vậy phần này là giao thức "
        "đánh giá có dữ liệu đã khóa, chưa phải bảng kết quả Precision, Recall, F1 hoặc confusion matrix.",
    )
    replace(find(doc, "Bảng 5.2. Ma trận quy trình"), "Bảng 5.2. Giao thức đánh giá extraction/matching trên pilot-v1", "Caption")
    replace(find_body(doc, "5.3. Đánh giá trải nghiệm"), "5.3. Giao thức đánh giá feedback và A/B", "Heading 2")
    replace(
        find(doc, "Hệ thống đã có cơ chế gán variant"),
        "Hệ thống có cơ chế gán variant theo session, lưu survey và tổng hợp Helpfulness, Actionability, No-answer-leak. Artifact readiness ngày 09/08/2026 "
        "xác nhận chưa có consent được phê duyệt và schema chưa lưu version/thời điểm consent. Vì vậy không được chạy A/B với người dùng thật hoặc suy ra variant thắng. "
        "Bảng 5.3 là checklist bắt buộc trước khi bắt đầu thu thập dữ liệu.",
    )
    replace(find(doc, "Bảng 5.3. Ma trận điều kiện"), "Bảng 5.3. Điều kiện mở khóa đánh giá A/B feedback", "Caption")
    replace(find_body(doc, "5.4. Thảo luận và nguy cơ"), "5.4. Threats to validity và giới hạn diễn giải", "Heading 2")
    replace(
        find(doc, "Kết quả kỹ thuật chịu ảnh hưởng"),
        "Validity nội tại chịu ảnh hưởng bởi provider AI, model, prompt, threshold, scenario và chất lượng annotation. Validity ngoại tại bị giới hạn vì pilot dùng transcript tổng hợp "
        "và chưa có mẫu người học hợp lệ. Validity kết luận bị giới hạn vì test unit/service không thay thế integration test PostgreSQL, E2E UI hoặc user study. "
        "Do đó, mọi kết luận của Chương 5 chỉ áp dụng cho artifact kỹ thuật và điều kiện chạy đã nêu.",
    )
    replace(find_body(doc, "5.5. Đối chiếu câu hỏi"), "5.5. Trả lời câu hỏi nghiên cứu theo mức bằng chứng", "Heading 2")
    replace(
        find(doc, "Nội dung trong bảng sau"),
        "Bảng 5.4 nối từng RQ với artifact cụ thể và giới hạn kết luận. Cách trình bày này không biến một implementation demo thành bằng chứng thực nghiệm về hiệu quả đào tạo.",
    )
    replace(find(doc, "Bảng 5.4. Đối chiếu"), "Bảng 5.4. Trả lời RQ1–RQ3 theo artifact hiện có", "Caption")


def improve_conclusion(doc):
    replace(
        find(doc, "Trong phạm vi phiên bản mã nguồn"),
        "Ở phạm vi source code và artifact đã đối chiếu, ReqSimulator hiện thực quản trị scenario có version, simulation session, transcript, controlled disclosure, "
        "structured extraction, matching, review/override và feedback survey. Các hạng mục được truy vết qua controller/service/entity, UML và log kiểm thử kỹ thuật. "
        "Đề tài không kết luận AI thay thế giảng viên, không kết luận độ chính xác matching và không kết luận hiệu quả feedback đối với người học.",
    )
    replace(
        find(doc, "Các hạn chế chính gồm"),
        "Giới hạn quan trọng nhất là evaluation chưa hoàn tất: annotation chưa dual review/adjudication, raw LLM holdout run chưa được phê duyệt và A/B chưa có consent. "
        "Ngoài ra, crawler chưa render đầy đủ SPA, chat chưa idempotent hoàn toàn, metadata evaluation chưa đầy đủ, video chưa upload end-to-end và coverage test frontend/integration DB còn mỏng.",
    )
    for prefix, value in (
        ("• Hoàn tất annotation", "• Hoàn tất dual annotation, adjudication, calibration và raw LLM holdout evaluation trước khi công bố metric extraction/matching."),
        ("• Bổ sung provenance", "• Lưu bất biến model, prompt, scoring policy, threshold, scenario version và dataset version cho mỗi evaluation."),
        ("• Bổ sung source URL", "• Cải thiện provenance scenario, topic classification và quality-level gating dựa trên corpus đã review."),
        ("• Hoàn thiện calibration", "• So sánh policy matching một–một/composite-aware trên dataset gán nhãn thay vì thay đổi threshold theo trực giác."),
        ("• Thực hiện user study", "• Chỉ thực hiện user study/A-B sau consent, kế hoạch lấy mẫu và phương pháp phân tích được phê duyệt."),
        ("• Bổ sung secure multipart", "• Hoàn thiện upload video an toàn, idempotency chat, contract test và E2E test với PostgreSQL cô lập."),
    ):
        replace(find(doc, prefix), value)


def apply_consistent_formatting(doc):
    for paragraph in doc.paragraphs:
        content = paragraph.text.strip()
        if not content:
            continue
        if content.startswith("Hình ") or content.startswith("Bảng "):
            paragraph.style = "Caption"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if paragraph.style.name == "Normal":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(0.75)
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.first_line_indent = Cm(0)

    # Table 0 is a hand-laid-out cover. Its blank paragraphs must not inherit
    # the body 1.5 line spacing or the cover splits across two pages.
    for paragraph in doc.tables[0].cell(0, 0).paragraphs:
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.first_line_indent = Cm(0)

    # Table 0 is a hand-laid-out cover. Other tables use the common grid style.
    for table in doc.tables[1:]:
        configure_table(table)

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(("CHƯƠNG ", "TÀI LIỆU THAM KHẢO", "PHỤ LỤC B.")):
            paragraph.paragraph_format.page_break_before = True


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    doc = Document(SOURCE)
    setup_styles(doc)
    improve_front_matter(doc)
    strengthen_chapter_two(doc)
    improve_chapter_five(doc)
    improve_conclusion(doc)
    apply_consistent_formatting(doc)
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
