from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SOURCE = Path(r"C:\Users\OS 11\Downloads\BaoCao_ReqSimulator_HUFLIT_IEEE_final_with_placeholder_validation_updated.docx")
TARGET = Path(r"C:\Users\OS 11\Downloads\BaoCao_ReqSimulator_HUFLIT_IEEE_final_research_rationale_updated.docx")


def text(p):
    return p.text.strip()


def find_nth(doc, prefix, occurrence=0):
    rows = [p for p in doc.paragraphs if text(p).startswith(prefix)]
    if len(rows) <= occurrence:
        raise RuntimeError(f"Missing {prefix} occurrence {occurrence}")
    return rows[occurrence]


def replace(p, value, italic=False, centered=False):
    p._p.clear_content()
    run = p.add_run(value)
    run.italic = italic
    run.font.size = Pt(11)
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def insert_after(anchor, value='', style=None):
    element = OxmlElement('w:p')
    anchor._p.addnext(element)
    from docx.text.paragraph import Paragraph
    p = Paragraph(element, anchor._parent)
    if style:
        p.style = style
    if value:
        p.add_run(value)
    return p


def shade(cell, color):
    tcpr = cell._tc.get_or_add_tcPr()
    node = OxmlElement('w:shd')
    node.set(qn('w:fill'), color)
    tcpr.append(node)


def set_cell(cell, value, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(value)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_after(doc, anchor, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True)
        shade(table.rows[0].cells[i], 'D9EAF7')
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Cm(width)
    anchor._p.addnext(table._tbl)
    return table


def main():
    doc = Document(SOURCE)
    h11 = find_nth(doc, '1.1. Lý do chọn đề tài', 1)
    p_reason_1 = find_nth(doc, 'Trong đào tạo phân tích', 0)
    p_reason_2 = find_nth(doc, 'ReqSimulator được định hướng', 0)

    replace(h11, '1.1. Vấn đề xuất phát và lý do chọn đề tài')
    replace(p_reason_1,
        'Đề tài bắt đầu từ việc người thực hiện tự xây dựng các kịch bản luyện phỏng vấn yêu cầu và nhận thấy phần khó không nằm ở việc nhớ định nghĩa FR, NFR hay business rule. Khó hơn là chọn câu hỏi tiếp theo khi câu trả lời của stakeholder còn mơ hồ, phát hiện trường hợp ngoại lệ và ghi lại thông tin theo cấu trúc có thể kiểm tra. Một buổi role-play trong lớp thường không đủ thời gian để lặp lại cùng tình huống, còn việc luyện với chatbot tự do lại khó xác định câu trả lời nào đã chạm đúng yêu cầu của kịch bản.')
    replace(p_reason_2,
        'Vì vậy, mục tiêu của ReqSimulator không phải tạo thêm một chatbot trả lời chung chung. Người thực hiện xây dựng một môi trường luyện tập có kịch bản phiên bản hóa, persona stakeholder, danh sách yêu cầu ẩn, rule gating và báo cáo sau phiên. Mỗi quyết định trong luồng này đều nhằm trả lời một câu hỏi thực hành: sinh viên đã hỏi được gì, còn bỏ sót gì và có thể hỏi tiếp theo theo hướng nào mà không lộ nguyên văn đáp án.')
    p_reason_3 = insert_after(p_reason_2, '', style=p_reason_2.style)
    replace(p_reason_3,
        'Phần nghiên cứu của đề tài được thể hiện bằng các artifact có thể đối chiếu: cấu hình scenario, transcript pilot, schema AAOC, UML, source code, test và dataset lock. Báo cáo không coi quan sát cá nhân này là khảo sát đại diện cho tất cả sinh viên; nó là lý do hình thành bài toán và phạm vi thiết kế của hệ thống.')

    h12_objective = find_nth(doc, '1.2. Mục tiêu nghiên cứu', 1)
    h12 = insert_after(p_reason_3, '1.2. Đối chiếu với công cụ hiện có và điểm khác biệt của ReqSimulator', style=h12_objective.style)
    comparison_intro = insert_after(h12, '', style=p_reason_1.style)
    replace(comparison_intro,
        'Bảng 1.1 là đối chiếu chức năng do người thực hiện thực hiện tại thời điểm phát triển đề tài; không phải benchmark hiệu năng hoặc đánh giá chất lượng thương mại của các sản phẩm được nêu.')
    caption_11 = insert_after(comparison_intro, '', style=p_reason_1.style)
    replace(caption_11, 'Bảng 1.1. Đối chiếu chức năng của ReqSimulator với các nhóm công cụ hiện có', italic=True, centered=True)
    table_after(doc, caption_11,
        ['Nhóm công cụ', 'Phù hợp cho', 'Khoảng trống với mục tiêu luyện elicitation', 'Điểm xử lý trong ReqSimulator'],
        [
            ('Chatbot LLM tổng quát, ví dụ ChatGPT', 'Trao đổi tự do, giải thích kiến thức, gợi ý câu hỏi.', 'Không có scenario do giảng viên kiểm soát, không có yêu cầu ẩn theo gate và không có cách chấm nhất quán theo từng phiên.', 'Persona gắn scenario; disclosure gating; transcript và matching với hidden requirement.'),
            ('Jira / Confluence và công cụ quản lý yêu cầu', 'Lưu tài liệu, ticket, backlog, liên kết công việc.', 'Hỗ trợ quản lý requirement sau khi đã có nội dung, không mô phỏng stakeholder để người học luyện khai thác thông tin.', 'Tập trung vào giai đoạn trước đặc tả: đặt câu hỏi, lấy bằng chứng, trích xuất và review.'),
            ('Quiz hoặc biểu mẫu tĩnh', 'Kiểm tra câu trả lời theo bộ câu hỏi cố định.', 'Không phản ánh việc câu hỏi tiếp theo phụ thuộc câu trả lời trước đó; khó mô hình hóa disclosure có điều kiện.', 'Hội thoại nhiều lượt, trạng thái persona và rule mở khóa requirement theo nội dung hỏi.'),
            ('ReqSimulator', 'Luyện một phiên elicitation và xem lại kết quả.', 'Không thay thế công cụ quản lý dự án hay khảo sát người học diện rộng.', 'Kết hợp scenario, controlled dialogue, structured extraction, one-to-one matching, review giảng viên và feedback an toàn.')
        ], [3.2, 4.0, 5.1, 5.2])

    # Shift Chapter 1 numbering after the new comparison section.
    renumbering = [
        ('1.2.1. Mục tiêu tổng quát', '1.3.1. Mục tiêu tổng quát'),
        ('1.2.2. Mục tiêu cụ thể', '1.3.2. Mục tiêu cụ thể'),
        ('1.3. Câu hỏi nghiên cứu', '1.4. Câu hỏi nghiên cứu'),
        ('1.4. Đối tượng và phạm vi nghiên cứu', '1.5. Đối tượng và phạm vi nghiên cứu'),
        ('1.5. Phương pháp nghiên cứu', '1.6. Phương pháp tự nghiên cứu và kiểm chứng'),
        ('1.6. Đóng góp của đề tài', '1.7. Đóng góp của đề tài'),
        ('1.7. Bố cục báo cáo', '1.8. Bố cục báo cáo'),
    ]
    objective = find_nth(doc, '1.3.1. Mục tiêu tổng quát', 0) if False else None
    # The current occurrences are body headings; rewriting in descending textual order avoids prefix collisions.
    for old, new in renumbering:
        heading = find_nth(doc, old, 1)
        replace(heading, new)
    replace(h12_objective, '1.3. Mục tiêu nghiên cứu')

    method_body = find_nth(doc, 'Đề tài kết hợp nghiên cứu tài liệu', 0)
    replace(method_body,
        'Người thực hiện kết hợp đọc tài liệu nền tảng với việc tự xây dựng và kiểm chứng artifact của hệ thống. Quy trình gồm: xác định các tình huống luyện tập từ scenario; mô hình hóa actor, use case, dữ liệu và luồng lỗi; hiện thực ba dịch vụ; viết test cho gating, retry, parsing JSON, matching và feedback; sau đó lưu log build/test. Dataset pilot-v1 được khóa bằng manifest gồm 10 transcript tổng hợp, checksum và split theo session. Vì annotation vẫn là version 1, chưa dual review/adjudication và raw LLM holdout run chưa được phê duyệt, báo cáo không dùng nó để khẳng định Precision, Recall, F1 hoặc hiệu quả tổng quát.')

    # Technology rationale follows the architecture, before the figure.
    h41 = find_nth(doc, '4.1. Kiến trúc tổng quan', 1)
    architecture_body = find_nth(doc, 'Hệ thống có ba dịch vụ chính', 0)
    h411 = insert_after(architecture_body, '4.1.1. Lý do lựa chọn công nghệ', style=h41.style)
    tech_intro = insert_after(h411, '', style=architecture_body.style)
    replace(tech_intro,
        'Các công nghệ được chọn theo ràng buộc của bài toán và source code hiện có, không dựa trên một benchmark tuyên bố công nghệ này tốt hơn mọi lựa chọn khác. Tiêu chí chính là tách trách nhiệm, kiểm soát contract giữa dịch vụ, khả năng kiểm thử và khả năng thay provider AI khi cần.')
    caption_41 = insert_after(tech_intro, '', style=architecture_body.style)
    replace(caption_41, 'Bảng 4.1. Lý do lựa chọn công nghệ trong ReqSimulator', italic=True, centered=True)
    table_after(doc, caption_41,
        ['Thành phần', 'Lựa chọn', 'Lý do trong phạm vi đề tài', 'Không khẳng định'],
        [
            ('Web client', 'Vite + TypeScript', 'SPA quy mô vừa, build nhanh và kiểm tra kiểu cho payload gọi API; phù hợp codebase hiện tại.', 'Không chứng minh tốt hơn React, Angular hay Vue ở mọi dự án.'),
            ('API nghiệp vụ', 'ASP.NET Core', 'Phù hợp backend C# hiện có, controller/DTO rõ ràng, phân quyền và persistence qua EF Core.', 'Không phải so sánh throughput với Node.js hoặc Spring.'),
            ('AI service', 'FastAPI (Python)', 'Tách gọi provider AI, parsing, retry, fallback, extraction và matching khỏi API nghiệp vụ; thuận tiện kiểm thử độc lập.', 'Không chứng minh FastAPI nhanh hơn mọi framework backend khác.'),
            ('Dữ liệu', 'PostgreSQL', 'Phù hợp quan hệ User–Session–Message–Evaluation và trường JSONB cho dữ liệu đánh giá/audit.', 'Không đánh giá benchmark so với MySQL hay SQLite.'),
            ('Provider LLM', 'Gemini 2.5 Flash mặc định, qua adapter đa provider', 'Dùng cho hội thoại/trích xuất trong cấu hình hiện tại; adapter cho phép thay provider và fallback khi lỗi.', 'Không kết luận model mặc định chính xác hoặc rẻ nhất nếu chưa benchmark có kiểm soát.')
        ], [2.7, 3.0, 6.9, 4.9])

    # Replace a generic limitation list with traceable implementation problems and resolutions.
    h47 = find_nth(doc, '4.7. Các giới hạn triển khai cần nêu rõ', 1)
    replace(h47, '4.7. Vấn đề trong quá trình thực hiện, cách xử lý và giới hạn còn lại')
    p47 = find_nth(doc, '• Chưa có bằng chứng nghiên cứu người dùng', 0)
    replace(p47,
        'Các vấn đề dưới đây xuất hiện trong quá trình hiện thực. Bảng 4.2 nêu nguyên nhân kỹ thuật, cách xử lý đã có và phần còn giới hạn; cách trình bày này giúp phân biệt việc đã sửa với việc mới dừng ở kế hoạch.')
    caption_42 = insert_after(p47, '', style=p47.style)
    replace(caption_42, 'Bảng 4.2. Vấn đề triển khai, nguyên nhân và cách xử lý', italic=True, centered=True)
    table_after(doc, caption_42,
        ['Vấn đề', 'Tại sao xảy ra', 'Cách xử lý / evidence', 'Giới hạn còn lại'],
        [
            ('20 test async bị fail', 'Môi trường test thiếu pytest-asyncio nên pytest không chạy coroutine có pytest.mark.asyncio.', 'Thêm requirements-dev.txt và cài pytest-asyncio==1.4.0; rerun đạt 92 passed, 11 warnings. Log ART-01.', 'Dependency test cần được commit/tag cùng source trước khi nộp.'),
            ('LLM có thể trả JSON sai format hoặc provider lỗi', 'Đầu ra mô hình là xác suất; API có thể rate-limit hoặc không trả JSON hợp lệ.', 'Pydantic validation, retry, parse repair và regex fallback; response có cờ isFallback; có test retry/parsing.', 'Fallback không thay thế output LLM trong đánh giá chính thức.'),
            ('Feedback có nguy cơ lộ hidden requirement', 'Nếu prompt đưa nguyên văn ground truth vào model, gợi ý có thể tiết lộ đáp án.', 'Variant B chỉ nhận ID/category/match score/AAOC; test xác nhận prompt không chứa SECRET_GROUND_TRUTH.', 'Cần user study có consent để đo no-answer-leak với người dùng thật.'),
            ('Matching dễ nhầm khi chỉ dựa từ giống nhau', 'Hai requirement có thể giống từ khóa nhưng khác actor, action, object hoặc điều kiện.', 'Lọc Type/Action/Object, AAOC 20/30/30/20, one-to-one assignment và lecturer override.', 'Threshold chưa được chốt bằng dual annotation + holdout raw LLM.'),
            ('A/B chưa thể kết luận', 'Chưa có consent được phê duyệt và schema chưa lưu consent version.', 'Ghi readiness gate, không chạy A/B/không điền winner.', 'Cần approval, consent versioning, cỡ mẫu và kế hoạch phân tích khóa trước thu thập.')
        ], [3.2, 4.3, 5.7, 4.3])
    # Remove the previous limitation-only text nodes, which are superseded by the table.
    for prefix in (
        'Matching hiện đã áp dụng AAOC',
        'Mô hình dữ liệu hiện đã tách Stakeholder',
        '• Video ingestion chưa phải upload multipart',
    ):
        paragraph = find_nth(doc, prefix, 0)
        paragraph._element.getparent().remove(paragraph._element)

    doc.save(TARGET)
    print(TARGET)


if __name__ == '__main__':
    main()
