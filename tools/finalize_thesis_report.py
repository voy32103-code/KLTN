from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SOURCE = Path(r"C:\Users\OS 11\Downloads\BaoCao_ReqSimulator_HUFLIT_IEEE_final_visual_tables.docx")
TARGET = Path(r"C:\Users\OS 11\Downloads\BaoCao_ReqSimulator_HUFLIT_IEEE_final_with_placeholder.docx")
ARTIFACT = Path(r"D:\KLTN\docs\audits_evaluation\verification_2026-08-09.md")

ANON_LOGIN = Path(r"C:\Users\OS 11\.codex\generated_images\019fe544-6430-7a11-86f6-cb8e3add5ff3\exec-a5d56d5a-3300-48b4-9e0f-75fa81891472.png")
ANON_LEADERBOARD = Path(r"C:\Users\OS 11\.codex\generated_images\019fe544-6430-7a11-86f6-cb8e3add5ff3\exec-d093372e-7d86-4026-8c83-305eb5283811.png")
ANON_CRUD = Path(r"C:\Users\OS 11\.codex\generated_images\019fe544-6430-7a11-86f6-cb8e3add5ff3\exec-1acd7eb0-bc29-464a-b751-2caf052a32a3.png")
ANON_REVIEW = Path(r"C:\Users\OS 11\.codex\generated_images\019fe544-6430-7a11-86f6-cb8e3add5ff3\exec-4a346193-8698-4483-828f-09e0d915df39.png")


def delete_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def delete_paragraph(paragraph):
    delete_element(paragraph._element)


def text_of(paragraph):
    return paragraph.text.strip()


def find_paragraph(doc, starts_with, occurrence=0):
    matches = [p for p in doc.paragraphs if text_of(p).startswith(starts_with)]
    if len(matches) <= occurrence:
        raise RuntimeError(f"Không tìm thấy đoạn: {starts_with}")
    return matches[occurrence]


def clear_paragraph(paragraph):
    paragraph._p.clear_content()


def set_paragraph(paragraph, value, bold=False, italic=False, align=None, size=None):
    clear_paragraph(paragraph)
    run = paragraph.add_run(value)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def copy_paragraph_format(source, target):
    target.style = source.style
    target.paragraph_format.left_indent = source.paragraph_format.left_indent
    target.paragraph_format.right_indent = source.paragraph_format.right_indent
    target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
    target.paragraph_format.space_before = source.paragraph_format.space_before
    target.paragraph_format.space_after = source.paragraph_format.space_after
    target.paragraph_format.line_spacing = source.paragraph_format.line_spacing


def insert_before(anchor, value='', style=None):
    new_p = OxmlElement('w:p')
    anchor._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    paragraph = Paragraph(new_p, anchor._parent)
    if style is not None:
        paragraph.style = style
    if value:
        paragraph.add_run(value)
    return paragraph


def insert_after(anchor, value='', style=None):
    new_p = OxmlElement('w:p')
    anchor._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    paragraph = Paragraph(new_p, anchor._parent)
    if style is not None:
        paragraph.style = style
    if value:
        paragraph.add_run(value)
    return paragraph


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_text(cell, value, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(value)
    r.bold = bold
    r.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def insert_table_before(doc, anchor, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        shade_cell(table.rows[0].cells[i], 'D9EAF7')
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    anchor._p.addprevious(table._tbl)
    return table


def insert_table_after(doc, anchor, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        shade_cell(table.rows[0].cells[i], 'D9EAF7')
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    anchor._p.addnext(table._tbl)
    return table


def table_contains(table, token):
    return token in ' '.join(cell.text for row in table.rows for cell in row.cells)


def replace_inline_image(doc, shape_index, image_path):
    shape = doc.inline_shapes[shape_index]
    blip = shape._inline.graphic.graphicData.pic.blipFill.blip
    relation_id = blip.embed
    part = doc.part.related_parts[relation_id]
    part._blob = image_path.read_bytes()


def body_paragraph(doc, text, template):
    p = insert_after(template, '')
    copy_paragraph_format(template, p)
    set_paragraph(p, text)
    return p


def set_heading(paragraph, text):
    set_paragraph(paragraph, text, bold=True)
    return paragraph


def remove_between(doc, first, last_exclusive):
    paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p._p == first._p)
    end = next(i for i, p in enumerate(paragraphs) if p._p == last_exclusive._p)
    for p in paragraphs[start + 1:end]:
        delete_paragraph(p)


def set_footer_note(doc):
    settings = doc.settings.element
    update = settings.find(qn('w:updateFields'))
    if update is None:
        update = OxmlElement('w:updateFields')
        settings.append(update)
    update.set(qn('w:val'), 'true')


def rebuild_cover(doc):
    """Build one clean cover page from the HUFLIT/FIT data supplied for this thesis."""
    acknowledgement = find_paragraph(doc, 'LỜI CẢM ƠN')
    for p in list(doc.paragraphs):
        if p._p == acknowledgement._p:
            break
        delete_paragraph(p)

    cover_lines = [
        ('BỘ GIÁO DỤC VÀ ĐÀO TẠO', 13, True, 0),
        ('TRƯỜNG ĐẠI HỌC NGOẠI NGỮ - TIN HỌC TP. HỒ CHÍ MINH', 13, True, 0),
        ('KHOA CÔNG NGHỆ THÔNG TIN', 13, True, 2),
        ('KHÓA LUẬN TỐT NGHIỆP', 16, True, 2),
        ('NGÀNH: CÔNG NGHỆ PHẦN MỀM', 13, True, 1),
        ('ỨNG DỤNG AI TRONG HỖ TRỢ ĐÀO TẠO NGHIỆP VỤ PHÂN TÍCH VÀ THIẾT KẾ PHẦN MỀM: THIẾT KẾ HỆ THỐNG MÔ PHỎNG TƯƠNG TÁC', 15, True, 3),
        ('Sinh viên thực hiện: VÕ HƯNG YÊN', 13, False, 0),
        ('MSSV: 23DH114176', 13, False, 0),
        ('Khóa: K29', 13, False, 0),
        ('Giảng viên hướng dẫn: ThS. ĐẶNG THỊ KIM GIAO', 13, False, 2),
        ('TP. Hồ Chí Minh, tháng 8 năm 2026', 13, False, 0),
    ]
    for value, size, bold, before in cover_lines:
        p = insert_before(acknowledgement, '')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before * 12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(value)
        r.bold = bold
        r.font.name = 'Times New Roman'
        r.font.size = Pt(size)
    page_break = insert_before(acknowledgement, '')
    page_break.add_run().add_break(WD_BREAK.PAGE)


def main():
    for required in [SOURCE, ARTIFACT, ANON_LOGIN, ANON_LEADERBOARD, ANON_CRUD, ANON_REVIEW]:
        if not required.exists():
            raise RuntimeError(f"Thiếu tệp cần thiết: {required}")

    doc = Document(SOURCE)

    # 1) Keep the original placeholder cover as requested, but remove the duplicate manual cover.
    placeholder_tables = [t for t in doc.tables if '[BỔ SUNG TÊN ĐỀ TÀI' in ' '.join(c.text for r in t.rows for c in r.cells)]
    if len(placeholder_tables) != 1:
        raise RuntimeError(f"Số bảng bìa placeholder bất thường: {len(placeholder_tables)}")
    placeholder_cover = placeholder_tables[0]
    acknowledgement = find_paragraph(doc, 'LỜI CẢM ƠN')
    body_children = list(doc._element.body)
    cover_index = body_children.index(placeholder_cover._tbl)
    acknowledgement_index = body_children.index(acknowledgement._p)
    for element in body_children[cover_index + 1:acknowledgement_index]:
        delete_element(element)

    # 2) Red editorial notes must not be present in a submission copy.
    for p in list(doc.paragraphs):
        if 'GHI CHÚ CẦN CHỈNH' in p.text or 'cần bổ sung trước khi nộp' in p.text.lower():
            delete_paragraph(p)

    # 3) Replace screenshots containing identifiable data with anonymized staging captures.
    # Shape order: 24 login, 27 leaderboard, 28 CRUD, 30 lecturer review.
    replace_inline_image(doc, 24, ANON_LOGIN)
    replace_inline_image(doc, 27, ANON_LEADERBOARD)
    replace_inline_image(doc, 28, ANON_CRUD)
    replace_inline_image(doc, 30, ANON_REVIEW)

    # 4) Strengthen the UI evidence note.
    ui_intro = find_paragraph(doc, 'Các hình dưới đây được chụp')
    set_paragraph(
        ui_intro,
        'Các hình giao diện trong mục này được chụp trên môi trường test/staging. Tên, email và định danh người dùng đã được ẩn danh; số liệu trên dashboard chỉ dùng để minh họa luồng vận hành, không phải số liệu thực nghiệm của Chương 5.'
    )

    # 5) Replace the whole theoretical foundation with evidence-focused, cited prose.
    ch2 = find_paragraph(doc, 'CHƯƠNG 2')
    ch3 = find_paragraph(doc, 'CHƯƠNG 3')
    template = find_paragraph(doc, '2.1.')
    body_template = find_paragraph(doc, 'Khai thác yêu cầu')
    remove_between(doc, ch2, ch3)

    cursor = ch2
    sections = [
        ('2.1. Khai thác yêu cầu phần mềm', [
            'Khai thác yêu cầu là hoạt động xác định, làm rõ và duy trì nhu cầu của các bên liên quan để tạo đầu vào có thể kiểm tra cho đặc tả hệ thống. ISO/IEC/IEEE 29148 yêu cầu yêu cầu phải rõ nguồn gốc, nhất quán và có thể truy vết; các tổng quan về kỹ nghệ yêu cầu cũng xem phỏng vấn, quan sát và phân tích tài liệu là những kỹ thuật trung tâm [1], [2]. Trong phạm vi ReqSimulator, phỏng vấn được chọn làm tình huống luyện tập vì người học phải chủ động đặt câu hỏi, xử lý câu trả lời không đầy đủ và ghi nhận bằng chứng từ hội thoại.',
            'Một buổi phỏng vấn có chất lượng không chỉ thu thập chức năng mong muốn mà còn làm rõ tác nhân, dữ liệu, quy tắc nghiệp vụ, ràng buộc và tiêu chí chấp nhận. Sommerville phân biệt yêu cầu người dùng với yêu cầu hệ thống; Zowghi và Coulin nhấn mạnh ảnh hưởng của giao tiếp, mâu thuẫn mục tiêu và tri thức ngầm của stakeholder [3], [4]. Vì vậy, hệ thống lưu hội thoại theo phiên và đối chiếu các phát biểu của người học với danh sách yêu cầu ẩn theo từng kịch bản, thay vì chấm bằng số lượng câu hỏi đơn thuần.'
        ]),
        ('2.2. Học tập theo kịch bản và phản hồi', [
            'Kịch bản tạo một bối cảnh công việc có giới hạn: lĩnh vực, vai trò stakeholder, mục tiêu nghiệp vụ và tập yêu cầu cần khai thác. Theo cách tiếp cận scenario-based design, người học có thể lặp lại một tình huống gần với thực tế, quan sát hệ quả của lựa chọn câu hỏi và so sánh các lần thực hiện trong cùng điều kiện [5]. Đây là cơ sở để ReqSimulator tổ chức kịch bản, persona stakeholder và danh sách yêu cầu ẩn thành các đối tượng phiên bản hóa.',
            'Phản hồi chỉ có ích khi chỉ ra chênh lệch giữa mục tiêu và kết quả hiện có, đồng thời gợi ý hành động tiếp theo có thể thực hiện. Hattie và Timperley mô tả phản hồi hiệu quả qua ba câu hỏi: mục tiêu là gì, người học đang ở đâu và bước kế tiếp nên làm gì [6]. Theo đó, báo cáo của ReqSimulator trình bày yêu cầu đã chạm tới, yêu cầu thiếu, mức độ khớp và gợi ý câu hỏi tiếp theo. Đây là lựa chọn thiết kế của hệ thống; tác động lên kết quả học tập cần được kiểm chứng bằng đánh giá người dùng riêng.'
        ]),
        ('2.3. Mô hình ngôn ngữ lớn trong hội thoại mô phỏng', [
            'Mô hình ngôn ngữ lớn có thể duy trì hội thoại theo vai, diễn đạt lại thông tin và hỗ trợ trích xuất cấu trúc từ văn bản tự do. Tuy nhiên, đầu ra của mô hình mang tính xác suất; các nghiên cứu đã chỉ ra rủi ro sinh nội dung không được hỗ trợ bởi nguồn và khó kiểm soát tính đúng đắn nếu thiếu ràng buộc [7], [8]. Do đó, ReqSimulator không dùng câu trả lời của mô hình như nguồn chân lý cho điểm số.',
            'Ở mức hiện thực, prompt được gắn với persona và kịch bản; dữ liệu đánh giá đi qua schema có cấu trúc, kiểm tra hợp lệ và các điều kiện chặn lỗi trước khi hiển thị. Danh sách yêu cầu ẩn vẫn là chuẩn đối chiếu, còn giảng viên có quyền review và điều chỉnh đánh giá. Các biện pháp này giảm rủi ro vận hành, nhưng không thay thế cho đánh giá độ chính xác trên dataset đã khóa.'
        ]),
        ('2.4. Truy vết và biểu diễn yêu cầu trích xuất', [
            'Khả năng truy vết giúp liên kết một yêu cầu với nguồn gốc, thay đổi và kết quả kiểm tra. Chuẩn 29148 yêu cầu duy trì liên kết giữa nhu cầu, yêu cầu và artifact kỹ thuật; các mô hình biểu diễn câu ở dạng vector có thể hỗ trợ tìm các phát biểu gần nghĩa [1], [9]. Trong ReqSimulator, mỗi mục trích xuất được liên kết với phiên hội thoại, câu nói nguồn, kịch bản và yêu cầu ẩn liên quan.',
            'Hệ thống chuẩn hóa phát biểu về các thành phần actor, action, object và constraint (AAOC) trước khi matching. Việc tách dữ liệu thô, kết quả chuẩn hóa và quyết định match giúp người review truy ngược vì sao một yêu cầu được tính điểm hoặc bị bỏ sót. Các phiên bản kịch bản và annotation được lưu cùng artifact đánh giá để kết quả có thể tái lập sau này.'
        ]),
        ('2.5. Matching ngữ nghĩa và quy tắc chấm', [
            'Matching ánh xạ một yêu cầu người học diễn đạt sang một yêu cầu ẩn phù hợp. Điểm tương đồng ngữ nghĩa chỉ được dùng để tạo ứng viên; quyết định cuối cùng cần kết hợp actor, hành động, đối tượng, ràng buộc và ngưỡng quy định bởi rubric. Các embedding dạng Sentence-BERT hỗ trợ so sánh ngữ nghĩa giữa câu, nhưng bản thân embedding không xác lập tiêu chí đúng/sai của miền nghiệp vụ [9].',
            'ReqSimulator phân loại kết quả thành exact, semantic, partial và missed. Exact yêu cầu các thành phần cốt lõi tương ứng; semantic cho phép cách diễn đạt khác nghĩa bề mặt nhưng không đổi nội dung; partial chỉ khớp một phần có ý nghĩa; missed là không tìm thấy bằng chứng đủ mạnh. Ngưỡng matching và rubric chỉ được chốt sau khi chạy dual annotation, adjudication và đo độ đồng thuận giữa người gắn nhãn, ví dụ bằng hệ số kappa của Cohen [10].'
        ]),
        ('2.6. Kiến trúc web, phân quyền và an toàn dữ liệu', [
            'ReqSimulator tách giao diện web, API nghiệp vụ, dịch vụ AI và cơ sở dữ liệu để mỗi lớp có trách nhiệm rõ ràng. Giao diện chỉ gửi yêu cầu đã xác thực; backend kiểm tra quyền, điều phối phiên và quản lý truy cập vào dữ liệu kịch bản. JWT là khuôn dạng compact để truyền claims và cần được kiểm tra chữ ký, thời hạn cùng ngữ cảnh sử dụng trước khi cấp quyền [11].',
            'Các yêu cầu an toàn trong hệ thống gồm kiểm soát truy cập theo vai trò, bảo vệ khóa dịch vụ AI, không đưa yêu cầu ẩn xuống client và ghi nhận sự kiện cần audit. Các nội dung này phù hợp với nhóm kiểm soát xác thực, phân quyền và logging trong OWASP ASVS [12]. Báo cáo mô tả hiện thực đã có; mức độ an toàn chỉ được kết luận sau kiểm thử bảo mật có phạm vi và artifact cụ thể.'
        ]),
    ]
    for heading, bodies in sections:
        cursor = insert_after(cursor, heading, style=template.style)
        set_heading(cursor, heading)
        for body in bodies:
            new_p = insert_after(cursor, '', style=body_template.style)
            copy_paragraph_format(body_template, new_p)
            set_paragraph(new_p, body)
            cursor = new_p

    # 6) Cite the UML basis at its first actual explanatory claim outside Chapter 2.
    uml_intro = find_paragraph(doc, 'Các sơ đồ UML')
    if '[13]' not in uml_intro.text:
        set_paragraph(uml_intro, uml_intro.text.rstrip() + ' [13]')

    # 7) Replace Bảng 5.1 with run artifacts rather than a broad "pass" claim.
    table_51 = next((t for t in doc.tables if table_contains(t, 'AI service unit tests') or table_contains(t, '54 tests pass')), None)
    if table_51 is None:
        raise RuntimeError('Không tìm thấy Bảng 5.1')
    delete_element(table_51._tbl)
    caption_51 = find_paragraph(doc, 'Bảng 5.1.')
    set_paragraph(caption_51, 'Bảng 5.1. Artifact kiểm chứng build/test tại snapshot 7efad62 (chạy ngày 09/08/2026)', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    insert_table_after(doc, caption_51,
        ['Hạng mục', 'Artifact / lệnh chạy', 'Kết quả và phạm vi chứng minh'],
        [
            ('Snapshot mã nguồn', 'Commit 7efad62, ngày 08/08/2026; log: docs/audits_evaluation/verification_2026-08-09.md', 'Mốc đối chiếu của lần chạy; chưa phải release tag.'),
            ('AI service unit tests', '.\\ai-service\\.venv\\Scripts\\python.exe -m pytest -q', '72 passed, 20 failed, 15 warning. Các lỗi async thiếu pytest-asyncio; không dùng kết quả này để tuyên bố test suite xanh.'),
            ('Backend API Release build', 'dotnet build .\\backend\\ReqSimulator.API\\ReqSimulator.API.csproj -c Release --no-restore', 'Pass; 0 warning, 0 error; 5.83 s.'),
            ('Backend integration build', 'dotnet build .\\backend\\ReqSimulator.API.IntegrationTests\\ReqSimulator.API.IntegrationTests.csproj -c Release --no-restore', 'Pass; 0 warning, 0 error; 2.44 s. Chưa chạy integration runtime với PostgreSQL cô lập.'),
            ('Frontend contract test', 'npm.cmd run test', 'Pass: 1 test kiểm tra payload override với DTO backend; không đại diện cho E2E.'),
            ('Frontend production build', 'npm.cmd run build', 'Pass: TypeScript + Vite 8.1.5; 15 modules; 513 ms.')
        ], [3.2, 7.4, 7.4])
    evidence_51 = find_paragraph(doc, 'Kiểm thử tập trung vào contract')
    set_paragraph(evidence_51, 'Các kiểm chứng kỹ thuật được đóng gói thành artifact để người đọc có thể tái chạy và đối chiếu. Bảng 5.1 ghi đúng kết quả tại snapshot đã nêu, bao gồm cả phần chưa đạt; log đầy đủ được lưu tại docs/audits_evaluation/verification_2026-08-09.md.')

    # 8) Add a defensible RQ-to-evidence mapping before Chapter 6.
    ch6 = find_paragraph(doc, 'CHƯƠNG 6')
    rq_heading = insert_before(ch6, '5.5. Đối chiếu câu hỏi nghiên cứu với bằng chứng', style=template.style)
    set_heading(rq_heading, '5.5. Đối chiếu câu hỏi nghiên cứu với bằng chứng')
    rq_intro = insert_before(ch6, 'Nội dung trong bảng sau phân biệt bằng chứng về thiết kế/hiện thực với bằng chứng thực nghiệm. Các kết luận không vượt quá phạm vi artifact đã có.')
    copy_paragraph_format(body_template, rq_intro)
    rq_caption = insert_before(ch6, 'Bảng 5.4. Đối chiếu RQ1–RQ3 với artifact và kết luận', style=caption_51.style)
    set_paragraph(rq_caption, 'Bảng 5.4. Đối chiếu RQ1–RQ3 với artifact và kết luận', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    insert_table_before(doc, ch6,
        ['Câu hỏi', 'Bằng chứng trong báo cáo', 'Kết luận có thể bảo vệ'],
        [
            ('RQ1: Có thể mô phỏng luyện phỏng vấn yêu cầu không?', 'Use case, activity/sequence, class diagram; ảnh giao diện staging; build API/frontend trong Bảng 5.1.', 'Đã chứng minh ở mức thiết kế và hiện thực luồng chọn kịch bản – hội thoại – review. Chưa suy ra mức cải thiện kỹ năng của người học.'),
            ('RQ2: Có thể đánh giá extraction/matching một cách truy vết không?', 'AAOC, hidden requirements, schema kết quả, log kiểm chứng; quy trình Bảng 5.2.', 'Pipeline và điểm truy vết đã được hiện thực. Chưa kết luận Precision/Recall/F1 cho toàn pipeline khi dataset và annotation chưa khóa.'),
            ('RQ3: Phản hồi có hữu ích hơn giữa các variant không?', 'Cơ chế feedback/variant trên giao diện; thiết kế đánh giá Bảng 5.3.', 'Mới có cơ chế và kế hoạch đo lường. Không có kết luận A/B khi chưa có consent, mẫu hợp lệ và phân tích định trước.')
        ], [4.1, 6.2, 7.7])

    # 9) Rebuild the IEEE reference list in order of first citation.
    refs_heading = find_paragraph(doc, 'TÀI LIỆU THAM KHẢO')
    appendix_heading = find_paragraph(doc, 'PHỤ LỤC')
    remove_between(doc, refs_heading, appendix_heading)
    references = [
        '[1] ISO/IEC/IEEE, ISO/IEC/IEEE 29148:2018, Systems and software engineering—Life cycle processes—Requirements engineering. Geneva, Switzerland: ISO, 2018.',
        '[2] B. Nuseibeh and S. Easterbrook, “Requirements engineering: A roadmap,” in Proc. Conf. on the Future of Software Engineering, 2000, pp. 35–46, doi: 10.1145/336512.336523.',
        '[3] I. Sommerville, Software Engineering, 10th ed. Boston, MA, USA: Pearson, 2015.',
        '[4] D. Zowghi and C. Coulin, “Requirements elicitation: A survey of techniques, approaches, and tools,” in Engineering and Managing Software Requirements, A. Aurum and C. Wohlin, Eds. Berlin, Germany: Springer, 2005, pp. 19–46.',
        '[5] J. M. Carroll, Scenario-Based Design: Envisioning Work and Technology in System Development. New York, NY, USA: Wiley, 1995.',
        '[6] J. Hattie and H. Timperley, “The power of feedback,” Review of Educational Research, vol. 77, no. 1, pp. 81–112, 2007, doi: 10.3102/003465430298487.',
        '[7] E. M. Bender, T. Gebru, A. McMillan-Major, and S. Shmitchell, “On the dangers of stochastic parrots: Can language models be too big?,” in Proc. ACM FAccT, 2021, pp. 610–623, doi: 10.1145/3442188.3445922.',
        '[8] Z. Ji et al., “Survey of hallucination in natural language generation,” ACM Comput. Surv., vol. 55, no. 12, pp. 1–38, 2023, doi: 10.1145/3571730.',
        '[9] N. Reimers and I. Gurevych, “Sentence-BERT: Sentence embeddings using Siamese BERT-networks,” in Proc. EMNLP-IJCNLP, 2019, pp. 3982–3992, doi: 10.18653/v1/D19-1410.',
        '[10] J. Cohen, “A coefficient of agreement for nominal scales,” Educational and Psychological Measurement, vol. 20, no. 1, pp. 37–46, 1960, doi: 10.1177/001316446002000104.',
        '[11] M. Jones, J. Bradley, and N. Sakimura, “JSON Web Token (JWT),” RFC 7519, May 2015, doi: 10.17487/RFC7519.',
        '[12] OWASP Foundation, OWASP Application Security Verification Standard, ver. 4.0.3, 2021. [Online]. Available: https://owasp.org/www-project-application-security-verification-standard/. [Accessed: Aug. 9, 2026].',
        '[13] Object Management Group, OMG Unified Modeling Language (OMG UML), ver. 2.5.1, Dec. 2017. [Online]. Available: https://www.omg.org/spec/UML/2.5.1/.'
    ]
    for ref in references:
        p = insert_before(appendix_heading, ref, style=body_template.style)
        copy_paragraph_format(body_template, p)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.paragraph_format.left_indent = Cm(0.8)

    # 10) Give Appendix A a concrete artifact manifest.
    appendix_b = find_paragraph(doc, 'PHỤ LỤC B')
    app_a_note = find_paragraph(doc, 'Bằng chứng kiểm thử gồm')
    set_paragraph(app_a_note, 'Phụ lục A liệt kê artifact kiểm chứng được dùng cho Bảng 5.1. Artifact lưu trong repository để người chấm có thể đối chiếu snapshot mã nguồn, lệnh chạy và giới hạn của từng kết quả.')
    app_a_caption = insert_before(appendix_b, 'Bảng A.1. Danh mục artifact kiểm chứng', style=caption_51.style)
    set_paragraph(app_a_caption, 'Bảng A.1. Danh mục artifact kiểm chứng', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    insert_table_before(doc, appendix_b,
        ['Mã artifact', 'Nội dung', 'Vị trí / trạng thái'],
        [
            ('ART-01', 'Log snapshot commit, môi trường, lệnh và output rút gọn.', 'docs/audits_evaluation/verification_2026-08-09.md — đã lưu.'),
            ('ART-02', 'AI service pytest run.', 'Ghi trong ART-01: 72 pass, 20 fail, 15 warning; cần bổ sung pytest-asyncio trước khi rerun.'),
            ('ART-03', 'Backend API và integration build Release.', 'Ghi trong ART-01: pass, 0 warning/error; integration runtime chưa chạy PostgreSQL cô lập.'),
            ('ART-04', 'Frontend contract test và production build.', 'Ghi trong ART-01: 1 test pass; build pass.')
        ], [2.2, 6.7, 8.9])

    # 11) Word requests a field update on opening; COM automation will also update it explicitly.
    set_footer_note(doc)
    doc.core_properties.comments = 'Bản nộp đã ẩn danh ảnh staging và liên kết artifact kiểm chứng.'
    doc.save(TARGET)
    print(TARGET)


if __name__ == '__main__':
    main()
