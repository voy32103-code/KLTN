"""Create a ReqSimulator thesis draft from the supplied Word sample.

The output intentionally keeps personal, institutional, reference, and
unmeasured-experiment information as visible placeholders for manual editing.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(r"C:\Users\OS 11\Downloads\BaoCao_KLTN.docx")
OUTPUT = Path(r"D:\KLTN\docs\BaoCao_ReqSimulator_draft.docx")


def set_run_font(run, size: float = 13, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")


def add_text(doc: Document, text: str, style: str = "Normal", *, bold: bool = False,
             center: bool = False, size: float = 13) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold)


def add_bullet(doc: Document, text: str) -> None:
    # The supplied sample does not define Word's built-in "List Bullet" style.
    # Keep bullets portable by using a normal paragraph with a literal marker.
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-9)
    paragraph.paragraph_format.space_after = Pt(3)
    set_run_font(paragraph.add_run(f"• {text}"), 13)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, 12, True)


def add_placeholder(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(f"[[{text}]]")
    set_run_font(run, 12, True)
    run.font.italic = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, 11, True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = ""
            run = cells[index].paragraphs[0].add_run(value)
            set_run_font(run, 10.5)
    doc.add_paragraph()


def add_toc_field(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Cập nhật mục lục trong Microsoft Word: chuột phải → Update Field."
    separate.append(placeholder)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])
    set_run_font(run, 12)


def clear_body_keep_cover(doc: Document) -> None:
    body = doc._element.body
    cover_table = doc.tables[0]._element if doc.tables else None
    for child in list(body):
        if child is cover_table or child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def write_cover(doc: Document) -> None:
    table = doc.tables[0]
    cell = table.cell(0, 0)
    cell.text = ""
    lines = [
        ("BỘ GIÁO DỤC VÀ ĐÀO TẠO", True, 13),
        ("TRƯỜNG ĐẠI HỌC NGOẠI NGỮ - TIN HỌC TP. HỒ CHÍ MINH", True, 13),
        ("KHOA CÔNG NGHỆ THÔNG TIN", True, 13),
        ("", False, 13),
        ("KHÓA LUẬN TỐT NGHIỆP", True, 18),
        ("NGÀNH: CÔNG NGHỆ PHẦN MỀM", True, 14),
        ("", False, 13),
        ("ĐỀ TÀI:", True, 14),
        ("[[TÊN ĐỀ TÀI CHÍNH THỨC]]", True, 16),
        ("", False, 13),
        ("Gợi ý: Xây dựng hệ thống mô phỏng phỏng vấn stakeholder hỗ trợ luyện tập khai thác yêu cầu phần mềm cho sinh viên.", False, 12),
        ("", False, 13),
        ("Sinh viên thực hiện: [[HỌ VÀ TÊN]]", False, 13),
        ("MSSV: [[MSSV]]", False, 13),
        ("Lớp: [[LỚP]]", False, 13),
        ("Giảng viên hướng dẫn: [[HỌ VÀ TÊN GIẢNG VIÊN]]", False, 13),
        ("", False, 13),
        ("TP. Hồ Chí Minh, tháng [[MM]] năm [[YYYY]]", False, 13),
    ]
    for text, bold, size in lines:
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_font(run, size, bold)


def add_chapter_1(doc: Document) -> None:
    add_text(doc, "CHƯƠNG 1 – TỔNG QUAN", "Heading 1", bold=True, center=True, size=15)
    add_text(doc, "1.1. Lý do chọn đề tài", "Heading 2", bold=True, size=14)
    add_text(doc, "Trong đào tạo phân tích và thiết kế phần mềm, sinh viên thường biết khái niệm yêu cầu phần mềm nhưng gặp khó khăn khi biến câu trả lời mơ hồ của khách hàng thành các câu hỏi làm rõ, câu hỏi ngoại lệ và quy tắc nghiệp vụ. Việc thực hành với stakeholder thật bị giới hạn bởi thời gian, chi phí và khả năng lặp lại. Vì vậy, cần một môi trường mô phỏng giúp người học luyện kỹ năng phỏng vấn requirement theo nhiều lượt hội thoại.")
    add_text(doc, "ReqSimulator được định hướng là hệ thống mô phỏng phỏng vấn stakeholder có scenario, persona, ground truth ẩn và cơ chế đánh giá sau phiên làm việc. Hệ thống không chỉ sinh hội thoại bằng mô hình ngôn ngữ mà còn kiểm soát thông tin được phép tiết lộ, lưu transcript, trích xuất requirement có cấu trúc, so khớp với ground truth và đưa ra phản hồi học tập.")
    add_text(doc, "1.2. Mục tiêu nghiên cứu", "Heading 2", bold=True, size=14)
    add_text(doc, "1.2.1. Mục tiêu tổng quát", "Heading 3", bold=True, size=13)
    add_text(doc, "Xây dựng và đánh giá ở mức kỹ thuật một hệ thống mô phỏng phỏng vấn stakeholder nhằm hỗ trợ sinh viên luyện tập khai thác yêu cầu phần mềm trong các scenario nghiệp vụ có kiểm soát.")
    add_text(doc, "1.2.2. Mục tiêu cụ thể", "Heading 3", bold=True, size=13)
    for item in [
        "Xây dựng luồng quản trị scenario gồm preview, chỉnh sửa, kiểm tra và publish version.",
        "Mô phỏng stakeholder có persona, trạng thái mood/patience và cơ chế disclosure theo rule gating.",
        "Trích xuất requirement từ transcript dưới dạng Actor–Action–Object–Condition, Type, Priority và Confidence.",
        "Chuẩn hóa, loại trùng, so khớp requirement với ground truth và tính coverage.",
        "Cung cấp feedback, gợi ý mô hình Use Case/ERD sơ bộ và màn hình review cho giảng viên.",
    ]:
        add_bullet(doc, item)
    add_text(doc, "1.3. Đối tượng và phạm vi nghiên cứu", "Heading 2", bold=True, size=14)
    add_text(doc, "Đối tượng nghiên cứu là quy trình elicitation requirement thông qua phỏng vấn stakeholder, mô phỏng scenario-based learning, structured extraction và đánh giá mức độ bao phủ requirement. Đối tượng sử dụng gồm sinh viên, giảng viên và quản trị viên.")
    add_text(doc, "Phạm vi triển khai hiện tại gồm frontend Vite/TypeScript, backend ASP.NET Core, AI service FastAPI và PostgreSQL. Scenario mẫu bao gồm đăng ký học phần, đặt lịch khám và quản lý tồn kho. Hệ thống không thay thế hoàn toàn hoạt động phỏng vấn thực tế, không đưa ra quyết định nghiệp vụ cho doanh nghiệp và chưa kết luận hiệu quả sư phạm khi chưa có nghiên cứu người dùng được chốt.")
    add_text(doc, "1.4. Phương pháp nghiên cứu", "Heading 2", bold=True, size=14)
    add_table(doc, ["Phương pháp", "Ứng dụng trong đề tài"], [
        ["Nghiên cứu lý thuyết", "Tổng hợp kiến thức về elicitation requirement, scenario-based learning, LLM, structured extraction và traceability."],
        ["Phân tích – thiết kế hệ thống", "Xác định actor, use case, kiến trúc ba dịch vụ, mô hình dữ liệu và luồng xử lý."],
        ["Thực nghiệm kỹ thuật", "Chạy build/test, dùng transcript pilot và kiểm tra các hành vi gating, extraction, matching."],
        ["Đánh giá người dùng", "[[BỔ SUNG SAU: mẫu, consent, survey, cách lấy mẫu và kết quả]]."],
    ])
    add_text(doc, "1.5. Đóng góp dự kiến", "Heading 2", bold=True, size=14)
    add_text(doc, "Đề tài đề xuất một môi trường luyện tập phỏng vấn requirement có ground truth ẩn, lưu vết transcript và phản hồi có thể truy vết. Đóng góp cần được diễn đạt theo bằng chứng triển khai và kết quả đánh giá thực tế; không khẳng định tăng hiệu quả học tập nếu chưa có số liệu nghiên cứu phù hợp.")
    add_text(doc, "1.6. Bố cục báo cáo", "Heading 2", bold=True, size=14)
    add_text(doc, "Báo cáo gồm sáu chương: Chương 1 giới thiệu đề tài; Chương 2 trình bày cơ sở lý thuyết; Chương 3 phân tích yêu cầu; Chương 4 mô tả thiết kế và triển khai; Chương 5 trình bày kiểm thử, đánh giá và giới hạn; Chương 6 kết luận và hướng phát triển.")


def add_chapter_2(doc: Document) -> None:
    doc.add_page_break()
    add_text(doc, "CHƯƠNG 2 – CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ LIÊN QUAN", "Heading 1", bold=True, center=True, size=15)
    add_text(doc, "2.1. Khai thác yêu cầu phần mềm và phỏng vấn stakeholder", "Heading 2", bold=True, size=14)
    add_text(doc, "Khai thác yêu cầu là hoạt động nhận diện, làm rõ, kiểm tra và thống nhất nhu cầu của các bên liên quan. Phỏng vấn stakeholder thường bắt đầu bằng câu hỏi mở để hiểu bối cảnh, sau đó dùng câu hỏi làm rõ, câu hỏi ràng buộc và câu hỏi ngoại lệ để phát hiện requirement chức năng, phi chức năng và quy tắc nghiệp vụ.")
    add_text(doc, "2.2. Học tập dựa trên scenario và phản hồi", "Heading 2", bold=True, size=14)
    add_text(doc, "Scenario-based learning tạo bối cảnh có mục tiêu, vai trò và dữ kiện giới hạn để người học thực hành. Trong ReqSimulator, scenario chứa hidden requirements và rule disclosure; điều này giúp đánh giá chất lượng câu hỏi thay vì chỉ đánh giá khả năng nhận lại đáp án từ chatbot.")
    add_text(doc, "2.3. Mô hình ngôn ngữ lớn trong môi trường học tập", "Heading 2", bold=True, size=14)
    add_text(doc, "Mô hình ngôn ngữ lớn hỗ trợ tạo phản hồi hội thoại và structured output, nhưng có nguy cơ hallucination, tiết lộ thông tin sai thời điểm và kết quả thiếu ổn định. Vì vậy, hệ thống cần prompt có ràng buộc, validation schema, gating, consistency checking, fallback an toàn và human review trước khi publish scenario.")
    add_text(doc, "2.4. Structured requirement extraction và normalization", "Heading 2", bold=True, size=14)
    add_text(doc, "Requirement được biểu diễn theo các thành phần Actor, Action, Object, Condition cùng Type, Priority và Confidence. Sau khi AI trả JSON, Pydantic kiểm tra schema; dữ liệu hợp lệ được chuẩn hóa alias, Unicode và khoảng trắng, sau đó loại trùng theo canonical key. Cách làm này giữ được cả bản gốc để review lẫn bản chuẩn hóa phục vụ đánh giá.")
    add_text(doc, "2.5. So khớp requirement và coverage", "Heading 2", bold=True, size=14)
    add_text(doc, "Hệ thống hiện tính similarity giữa requirement trích xuất và hidden requirement bằng embedding qua API, sau đó thực hiện gán một-một. Kết quả gồm exact, semantic, partial hoặc missed; coverage tính full match là 1 điểm và partial match là 0,5 điểm. Ngưỡng là policy cần được calibration, không phải kết luận khoa học tuyệt đối.")
    add_text(doc, "2.6. Kiến trúc web và an toàn dữ liệu", "Heading 2", bold=True, size=14)
    add_text(doc, "Kiến trúc triển khai tách frontend, backend công khai và AI service nội bộ. Backend quản lý JWT, phân quyền, session, versioning và lưu PostgreSQL. AI service thực hiện chat, gating, extraction, evaluation, crawler và video processing. Các ranh giới này giúp hạn chế việc client truy cập trực tiếp provider AI hoặc hidden requirement.")
    add_placeholder(doc, "BỔ SUNG TRÍCH DẪN HỌC THUẬT CHO CÁC MỤC 2.1–2.6 THEO CHUẨN TRÍCH DẪN CỦA TRƯỜNG")


def add_chapter_3(doc: Document) -> None:
    doc.add_page_break()
    add_text(doc, "CHƯƠNG 3 – PHÂN TÍCH YÊU CẦU HỆ THỐNG", "Heading 1", bold=True, center=True, size=15)
    add_text(doc, "3.1. Bối cảnh và các bên liên quan", "Heading 2", bold=True, size=14)
    add_text(doc, "ReqSimulator phục vụ ba nhóm người dùng: sinh viên thực hiện phỏng vấn mô phỏng; giảng viên rà soát transcript và kết quả; quản trị viên quản lý tài khoản và scenario. AI provider là thành phần kỹ thuật nội bộ, không phải actor nghiệp vụ trong Use Case Diagram.")
    add_caption(doc, "Bảng 3.1. Actor và trách nhiệm chính")
    add_table(doc, ["Actor", "Trách nhiệm"], [
        ["Sinh viên", "Chọn scenario/persona, chat, kết thúc phiên và xem feedback."],
        ["Giảng viên", "Rà soát session, hidden requirement, matching và override kết quả khi cần."],
        ["Quản trị viên", "Quản lý người dùng, tạo preview scenario từ URL/video, chỉnh sửa và publish."],
    ])
    add_text(doc, "3.2. Yêu cầu chức năng", "Heading 2", bold=True, size=14)
    add_table(doc, ["Mã", "Yêu cầu chức năng"], [
        ["FR-01", "Hệ thống cho phép người dùng đăng ký và đăng nhập theo vai trò."],
        ["FR-02", "Sinh viên chọn scenario, persona và khởi tạo session phỏng vấn."],
        ["FR-03", "Hệ thống lưu message, question type và persona state theo session."],
        ["FR-04", "AI trả lời stakeholder theo rule gating và consistency check."],
        ["FR-05", "Khi kết thúc session, hệ thống trích xuất, chuẩn hóa và đánh giá requirement."],
        ["FR-06", "Giảng viên xem transcript, hidden requirement, matching và override kết quả."],
        ["FR-07", "Admin tạo preview, chỉnh sửa và publish scenario theo version."],
    ])
    add_text(doc, "3.3. Yêu cầu phi chức năng và quy tắc nghiệp vụ", "Heading 2", bold=True, size=14)
    add_table(doc, ["Nhóm", "Nội dung"], [
        ["Bảo mật", "JWT, phân quyền role, ownership check, rate limiting và AI internal key."],
        ["Tính nhất quán", "Session finalization sử dụng lease và unique evaluation để giảm kết quả trùng."],
        ["Khả năng truy vết", "Scenario version, transcript, extracted requirement, match detail và lecturer override được lưu."],
        ["Nghiệp vụ", "Hidden requirement không được gửi trực tiếp cho client sinh viên trước khi được disclosure hợp lệ."],
    ])
    add_text(doc, "3.4. Use case và traceability", "Heading 2", bold=True, size=14)
    add_text(doc, "Use Case Diagram, Activity Diagram và Sequence Diagram được đặt trong tài liệu Mermaid của dự án. Khi hoàn thiện bản nộp, cần render các sơ đồ này thành hình, đánh số và chèn vào đúng vị trí trong báo cáo.")
    add_placeholder(doc, "CHÈN HÌNH 3.1 – USE CASE DIAGRAM")
    add_placeholder(doc, "CHÈN BẢNG TRUY VẾT USE CASE → MODULE → TEST CASE")


def add_chapter_4(doc: Document) -> None:
    doc.add_page_break()
    add_text(doc, "CHƯƠNG 4 – THIẾT KẾ VÀ TRIỂN KHAI HỆ THỐNG", "Heading 1", bold=True, center=True, size=15)
    add_text(doc, "4.1. Kiến trúc tổng quan", "Heading 2", bold=True, size=14)
    add_text(doc, "Hệ thống có ba dịch vụ chính. Frontend Vite/TypeScript cung cấp giao diện cho student, lecturer và admin. Backend ASP.NET Core là ranh giới công khai, xử lý xác thực, phân quyền, persistence và orchestration. AI service FastAPI xử lý persona chat, gating, structured extraction, evaluation, crawler và video. PostgreSQL lưu dữ liệu nghiệp vụ và lịch sử phiên.")
    add_placeholder(doc, "CHÈN HÌNH 4.1 – KIẾN TRÚC TỔNG QUAN (RENDER TỪ docs/architecture_specs/reqsimulator_mermaid_uml_diagrams.md)")
    add_text(doc, "4.2. Thiết kế dữ liệu", "Heading 2", bold=True, size=14)
    add_text(doc, "Các thực thể chính gồm User, Scenario, Persona, HiddenRequirement, SimulationSession, Message, ExtractedRequirement, EvaluationResult và RequirementMatch. Session liên kết với đúng scenario version và persona tại thời điểm bắt đầu, nhờ đó các phiên lịch sử vẫn có ngữ cảnh ổn định khi scenario mới được publish.")
    add_placeholder(doc, "CHÈN HÌNH 4.2 – CLASS DIAGRAM")
    add_text(doc, "4.3. Luồng hội thoại có kiểm soát", "Heading 2", bold=True, size=14)
    add_text(doc, "Khi sinh viên gửi câu hỏi, backend đọc history và persona state rồi gọi AI service. AI service phân loại câu hỏi, xác định requirement có thể tiết lộ, đưa đúng tập dữ kiện cho prompt stakeholder và kiểm tra consistency của câu trả lời. Backend chỉ lưu state update cần thiết; nội dung hidden requirement vừa mở khóa không được trả trực tiếp trong response client.")
    add_placeholder(doc, "CHÈN HÌNH 4.3 – SEQUENCE DIAGRAM CỦA PHIÊN PHỎNG VẤN")
    add_text(doc, "4.4. Pipeline extraction và evaluation", "Heading 2", bold=True, size=14)
    add_text(doc, "Khi session kết thúc, AI service nhận transcript, gọi prompt extraction để sinh requirement JSON có cấu trúc, kiểm tra bằng Pydantic và retry khi dữ liệu không hợp lệ. Requirement được chuẩn hóa, loại trùng và lưu cả raw/normalized data. Evaluation dùng similarity embedding, candidate threshold và gán một-một để tính coverage cùng feedback.")
    add_placeholder(doc, "CHÈN HÌNH 4.4 – ACTIVITY DIAGRAM EXTRACTION, NORMALIZATION VÀ EVALUATION")
    add_text(doc, "4.5. Quản trị scenario và video/URL ingestion", "Heading 2", bold=True, size=14)
    add_text(doc, "Admin có thể tạo scenario draft từ URL hoặc video/audio. Luồng UI chính tạo preview trước khi publish để admin chỉnh sửa context, hidden requirement, gate và điều kiện disclosure. Video hiện nhận đường dẫn file cục bộ tại máy chạy AI service, có thể tách audio bằng FFmpeg rồi gửi Gemini File API; secure browser upload là hướng phát triển tiếp theo.")
    add_text(doc, "4.6. Các giới hạn triển khai cần nêu rõ", "Heading 2", bold=True, size=14)
    add_bullet(doc, "Chưa có bằng chứng nghiên cứu người dùng đủ để kết luận hệ thống cải thiện kỹ năng elicitation.")
    add_bullet(doc, "Matching hiện dựa trên similarity một-một và threshold; chưa phải mô hình trọng số AAOC 20/30/30/20 với hard filter đầy đủ.")
    add_bullet(doc, "Mô hình dữ liệu hiện chưa tách Stakeholder khỏi Persona để tái sử dụng hoàn toàn.")
    add_bullet(doc, "Video ingestion chưa phải upload multipart từ trình duyệt và cần kiểm tra đường dẫn ở môi trường triển khai.")


def add_chapter_5(doc: Document) -> None:
    doc.add_page_break()
    add_text(doc, "CHƯƠNG 5 – KIỂM THỬ, ĐÁNH GIÁ VÀ THẢO LUẬN", "Heading 1", bold=True, center=True, size=15)
    add_text(doc, "5.1. Mục tiêu kiểm thử", "Heading 2", bold=True, size=14)
    add_text(doc, "Kiểm thử tập trung vào contract giữa các dịch vụ, gating requirement, fallback, structured extraction, normalization, matching, coverage, phân quyền và build của ba thành phần. Test pass chỉ chứng minh hành vi trên các test case đã có; không thay thế đánh giá thực nghiệm với người học.")
    add_caption(doc, "Bảng 5.1. Kết quả kiểm chứng kỹ thuật đã ghi nhận")
    add_table(doc, ["Hạng mục", "Kết quả ghi nhận", "Diễn giải"], [
        ["AI service unit tests", "54 tests pass", "Bao gồm extraction, validation, normalization, matching, evaluation, gating và fallback."],
        ["Backend Release build", "Pass; 0 warning, 0 error", "Kiểm tra biên dịch backend ASP.NET Core."],
        ["Backend integration project build", "Pass; 0 warning, 0 error", "Chỉ build; chưa chạy runtime vì cần PostgreSQL test cô lập."],
        ["Frontend test", "Pass", "Hiện là contract test; chưa đủ chứng minh toàn bộ UI."],
        ["Frontend production build", "Pass", "Vite/TypeScript build thành công."],
    ])
    add_text(doc, "5.2. Đánh giá pipeline extraction và matching", "Heading 2", bold=True, size=14)
    add_text(doc, "Dự án có pilot dataset gồm transcript và annotation. Bộ dữ liệu này có thể được dùng để đánh giá Precision, Recall, F1 và độ ổn định threshold sau khi annotation được khóa, quy trình lấy mẫu được xác nhận và script tái lập được chạy. Bản nháp này không tự điền các tỷ lệ chính xác khi chưa có kết quả đã xác minh.")
    add_placeholder(doc, "CHÈN BẢNG 5.2 – KẾT QUẢ EXTRACTION/MATCHING TRÊN DATASET ĐÃ KHÓA")
    add_text(doc, "5.3. Đánh giá trải nghiệm feedback", "Heading 2", bold=True, size=14)
    add_text(doc, "Feedback hiện dựa trên match detail, requirement missed/partial và extraction cần rà soát. Để kết luận về tính hữu ích, cần triển khai khảo sát có consent, xác định mẫu hợp lệ, variant feedback, tiêu chí helpfulness/actionability/no-answer-leak và phương pháp xử lý dữ liệu thiếu.")
    add_placeholder(doc, "CHÈN BẢNG 5.3 – KẾT QUẢ KHẢO SÁT NGƯỜI DÙNG SAU KHI THU THẬP ĐỦ MẪU")
    add_text(doc, "5.4. Thảo luận và nguy cơ ảnh hưởng đến validity", "Heading 2", bold=True, size=14)
    add_text(doc, "Kết quả kỹ thuật chịu ảnh hưởng bởi provider AI, prompt, model embedding, scenario, threshold và chất lượng annotation. Những yếu tố này cần được phiên bản hóa khi chạy thực nghiệm. Ngoài ra, một bộ test nhỏ hoặc demo thành công không đủ để suy luận khả năng vận hành production hay hiệu quả học tập dài hạn.")


def add_chapter_6(doc: Document) -> None:
    doc.add_page_break()
    add_text(doc, "CHƯƠNG 6 – KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "Heading 1", bold=True, center=True, size=15)
    add_text(doc, "6.1. Kết quả đạt được", "Heading 2", bold=True, size=14)
    add_text(doc, "Đề tài đã xây dựng được nền tảng ReqSimulator gồm quản lý scenario version, mô phỏng stakeholder theo persona, hội thoại có gating, structured extraction, normalization, matching một-một, coverage, feedback, lecturer review và luồng preview/edit/publish. Hệ thống được tổ chức theo kiến trúc frontend–backend–AI service và có các kiểm tra build/test ở mức kỹ thuật.")
    add_text(doc, "6.2. Hạn chế", "Heading 2", bold=True, size=14)
    add_text(doc, "Các hạn chế chính gồm thiếu đánh giá người dùng được chốt, độ phụ thuộc vào provider AI, giới hạn của matching hiện tại, chưa tách hoàn toàn Stakeholder/Persona, provenance nguồn scenario chưa đầy đủ và video ingestion chưa hỗ trợ upload browser an toàn.")
    add_text(doc, "6.3. Hướng phát triển", "Heading 2", bold=True, size=14)
    for item in [
        "Hoàn tất annotation, calibration và holdout evaluation cho extraction/matching.",
        "Tách Stakeholder và Persona; mở rộng nhiều stakeholder/persona cho từng scenario.",
        "Bổ sung source URL/provenance, topic classification và quality-level gating.",
        "Triển khai matching AAOC trọng số có candidate filtering theo Type/Action/Object.",
        "Thực hiện user study, survey và phân tích định lượng/định tính có thể tái lập.",
        "Bổ sung secure multipart video upload, idempotency chat và UI test end-to-end.",
    ]:
        add_bullet(doc, item)


def add_references_and_appendices(doc: Document) -> None:
    doc.add_page_break()
    add_text(doc, "TÀI LIỆU THAM KHẢO", "Heading 1", bold=True, center=True, size=15)
    add_placeholder(doc, "BỔ SUNG DANH MỤC TÀI LIỆU HỌC THUẬT ĐÃ KIỂM CHỨNG THEO CHUẨN IEEE/APA CỦA TRƯỜNG")
    add_text(doc, "Tài liệu nội bộ dùng làm evidence khi viết bản nháp:")
    for item in [
        "docs/README.md — trạng thái codebase và quy ước tài liệu.",
        "docs/research_academic/master_project_summary.md — tóm tắt phạm vi triển khai.",
        "docs/IMPLEMENTATION-READINESS.md — các hạng mục hoàn thiện và giới hạn.",
        "docs/architecture_specs/reqsimulator_mermaid_uml_diagrams.md — mã sơ đồ Mermaid.",
        "docs/research_academic/pilot_dataset/ — transcript và annotation pilot.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()
    add_text(doc, "PHỤ LỤC", "Heading 1", bold=True, center=True, size=15)
    add_text(doc, "Phụ lục A. Danh sách sơ đồ", "Heading 2", bold=True, size=14)
    add_text(doc, "Chèn bản render Use Case, Activity, Sequence, State, Class, kiến trúc tổng quan và quy trình từ tài liệu Mermaid. Mỗi hình cần có số thứ tự, tiêu đề và nguồn.")
    add_text(doc, "Phụ lục B. Hướng dẫn demo", "Heading 2", bold=True, size=14)
    add_text(doc, "Luồng demo đề xuất: đăng nhập → chọn scenario/persona → đặt câu hỏi theo nhiều lượt → kết thúc session → xem extraction/evaluation → lecturer review → admin preview/edit/publish scenario.")
    add_text(doc, "Phụ lục C. Bằng chứng kiểm thử", "Heading 2", bold=True, size=14)
    add_text(doc, "Đính kèm log build/test, phiên bản code, cấu hình môi trường không chứa secret, dataset version, annotation guideline, survey/consent và output evaluation sau khi các artifact này được khóa.")
    add_internal_design_appendix(doc)


def clean_markdown_text(text: str) -> str:
    """Keep internal specification content editable in Word without Markdown noise."""
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = text.replace("|", " — ")
    return re.sub(r"\s+", " ", text).strip(" -")


def add_internal_design_appendix(doc: Document) -> None:
    """Append the long design document as a clearly-labelled, editable technical appendix.

    It intentionally remains an internal design artifact, rather than being presented
    as experimental evidence or functionality proven by the current codebase.
    """
    source = Path(r"D:\KLTN\docs\architecture_specs\REQSIMULATOR_DUAL_LOOP_ADAPTIVE_DESIGN.md")
    if not source.is_file():
        return
    doc.add_page_break()
    add_text(doc, "PHỤ LỤC D. ĐẶC TẢ THIẾT KẾ KỸ THUẬT MỞ RỘNG", "Heading 1", bold=True, center=True, size=15)
    add_text(doc, "Phụ lục này chuyển thể tài liệu thiết kế nội bộ để làm dày phần mô tả kỹ thuật của khóa luận. Nội dung ở đây là đặc tả/định hướng cần được người thực hiện đối chiếu với mã nguồn và dữ liệu thực nghiệm trước khi nộp. Nó không được dùng để khẳng định một tính năng đã triển khai hay một kết quả nghiên cứu đã được kiểm chứng.", bold=True)
    add_placeholder(doc, "RÀ SOÁT TỪNG MỤC DƯỚI ĐÂY: GIỮ NẾU CÓ EVIDENCE, HOẶC CHUYỂN SANG HƯỚNG PHÁT TRIỂN")

    in_code = False
    for raw in source.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("<!--"):
            continue
        if line.startswith("```") or line.startswith("~~~"):
            in_code = not in_code
            continue
        if in_code:
            add_text(doc, line, size=10.5)
            continue
        if line.startswith("# "):
            add_text(doc, clean_markdown_text(line[2:]), "Heading 2", bold=True, size=14)
        elif line.startswith("## "):
            add_text(doc, clean_markdown_text(line[3:]), "Heading 3", bold=True, size=13)
        elif line.startswith("### "):
            add_text(doc, clean_markdown_text(line[4:]), bold=True, size=13)
        elif line.startswith(("- ", "* ")):
            add_bullet(doc, clean_markdown_text(line[2:]))
        elif re.match(r"^\d+\.\s+", line):
            add_text(doc, clean_markdown_text(line), size=13)
        elif line.startswith("|"):
            add_text(doc, clean_markdown_text(line), size=11)
        else:
            add_text(doc, clean_markdown_text(line), size=13)


def main() -> None:
    if not SOURCE.is_file():
        raise FileNotFoundError(f"Sample report not found: {SOURCE}")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)
    clear_body_keep_cover(doc)
    write_cover(doc)
    doc.add_page_break()
    add_text(doc, "LỜI CẢM ƠN", "Heading 1", bold=True, center=True, size=15)
    add_text(doc, "Người thực hiện xin trân trọng cảm ơn [[HỌ VÀ TÊN GIẢNG VIÊN HƯỚNG DẪN]], quý thầy cô Khoa Công nghệ Thông tin và các cá nhân đã hỗ trợ trong quá trình thực hiện khóa luận. [[CHỈNH SỬA LỜI CẢM ƠN THEO THÔNG TIN THỰC TẾ]].")
    doc.add_page_break()
    add_text(doc, "MỤC LỤC", "Heading 1", bold=True, center=True, size=15)
    add_toc_field(doc)
    doc.add_page_break()
    add_text(doc, "DANH MỤC CÁC KÝ HIỆU, CHỮ VIẾT TẮT VÀ THUẬT NGỮ", "Heading 1", bold=True, center=True, size=15)
    add_table(doc, ["STT", "Từ viết tắt", "Viết đầy đủ", "Ý nghĩa"], [
        ["1", "AI", "Artificial Intelligence", "Trí tuệ nhân tạo"],
        ["2", "LLM", "Large Language Model", "Mô hình ngôn ngữ lớn"],
        ["3", "FR", "Functional Requirement", "Yêu cầu chức năng"],
        ["4", "NFR", "Non-functional Requirement", "Yêu cầu phi chức năng"],
        ["5", "BR", "Business Rule", "Quy tắc nghiệp vụ"],
        ["6", "API", "Application Programming Interface", "Giao diện lập trình ứng dụng"],
        ["7", "JWT", "JSON Web Token", "Cơ chế token xác thực"],
        ["8", "UML", "Unified Modeling Language", "Ngôn ngữ mô hình hóa thống nhất"],
    ])
    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc)
    add_chapter_4(doc)
    add_chapter_5(doc)
    add_chapter_6(doc)
    add_references_and_appendices(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
