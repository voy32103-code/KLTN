from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(r"D:\KLTN\KLTN_ReqSimulator_mau_don_gian.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_toc(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Cập nhật mục lục trong Word bằng Ctrl+A, F9"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Trang | ").font.name = "Times New Roman"
    add_field(p, "PAGE")


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(6)
    normal.paragraph_format.space_after = Pt(8)

    for name, size, bold, align in [
        ("Heading 1", 15, True, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 14, True, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 13, True, WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.alignment = align
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.keep_with_next = True

    if "Caption Custom" not in [s.name for s in doc.styles]:
        caption = doc.styles.add_style("Caption Custom", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = doc.styles["Caption Custom"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(11)
    caption.font.italic = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
    return table


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption Custom")
    p.add_run(text)


def add_body(doc, text):
    doc.add_paragraph(text)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27)
    p.add_run(text)


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3)
    sec.bottom_margin = Cm(3)
    sec.left_margin = Cm(3.5)
    sec.right_margin = Cm(2)
    add_page_number(sec)
    configure_styles(doc)

    # Cover
    for text, size, bold in [
        ("TRƯỜNG ĐẠI HỌC NGOẠI NGỮ - TIN HỌC TP.HCM", 14, False),
        ("KHOA CÔNG NGHỆ THÔNG TIN", 14, False),
        ("", 8, False),
        ("KHÓA LUẬN TỐT NGHIỆP", 18, True),
        ("HỆ THỐNG MÔ PHỎNG PHỎNG VẤN STAKEHOLDER ĐỂ LUYỆN KỸ NĂNG KHAI THÁC YÊU CẦU", 20, True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold
    for _ in range(4):
        doc.add_paragraph()
    for text in [
        "GIẢNG VIÊN HƯỚNG DẪN: ........................................",
        "SINH VIÊN THỰC HIỆN: ........................................",
        "MSSV: ........................................................",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text).font.size = Pt(14)
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TP. HỒ CHÍ MINH, THÁNG 8 NĂM 2026")
    r.bold = True
    r.font.size = Pt(14)

    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LỜI CẢM ƠN")
    r.bold = True
    r.font.size = Pt(15)
    add_body(doc, "Nhóm thực hiện xin chân thành cảm ơn giảng viên hướng dẫn và Khoa Công nghệ Thông tin đã hỗ trợ trong quá trình hình thành ý tưởng, xây dựng hệ thống và hoàn thiện báo cáo thử nghiệm này.")
    add_body(doc, "Bản mẫu dưới đây chỉ minh họa cách tổ chức nội dung và trình bày; các số liệu, tên người và thông tin hành chính sẽ được cập nhật trong bản chính thức.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LỜI CAM ĐOAN")
    r.bold = True
    r.font.size = Pt(15)
    add_body(doc, "Nhóm thực hiện cam đoan nội dung trong báo cáo chính thức sẽ được đối chiếu với mã nguồn, tài liệu và kết quả kiểm thử thực tế. Những phần chưa có dữ liệu thực nghiệm sẽ được trình bày là giới hạn hoặc hướng phát triển.")

    doc.add_heading("BẢNG TỪ VIẾT TẮT", level=1)
    add_table(doc, ["STT", "Từ viết tắt", "Mô tả"], [["1", "AI", "Artificial Intelligence"], ["2", "LLM", "Large Language Model"], ["3", "BA", "Business Analyst"], ["4", "NFR", "Non-functional Requirement"]])

    doc.add_heading("MỤC LỤC", level=1)
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    # Chapter 1
    doc.add_heading("CHƯƠNG 1. GIỚI THIỆU", level=1)
    doc.add_heading("1.1. Lý do chọn đề tài", level=2)
    add_body(doc, "Nhiều sinh viên có thể ghi nhớ khái niệm requirement nhưng gặp khó khăn khi phải đặt câu hỏi nối tiếp với một stakeholder nói mơ hồ. ReqSimulator được đề xuất như một môi trường luyện tập lặp lại, trong đó sinh viên đóng vai Business Analyst và phỏng vấn stakeholder ảo.")
    doc.add_heading("1.2. Mục tiêu và phạm vi", level=2)
    add_bullet(doc, "Xây dựng phiên phỏng vấn nhiều lượt có persona và hidden requirement.")
    add_bullet(doc, "Lưu transcript và tạo phản hồi ở mức requirement.")
    add_bullet(doc, "Giữ nguyên phiên bản scenario để kết quả cũ có thể truy vết.")
    doc.add_heading("1.3. Phương pháp thực hiện", level=2)
    add_body(doc, "Đề tài kết hợp phân tích yêu cầu, thiết kế kiến trúc, phát triển phần mềm theo module và kiểm thử theo từng tầng. Các tuyên bố về hiệu quả giáo dục chỉ được đưa ra sau khi có pilot với protocol và dữ liệu tái lập.")

    # Chapter 2
    doc.add_heading("CHƯƠNG 2. CƠ SỞ LÝ THUYẾT", level=1)
    doc.add_heading("2.1. Khai thác yêu cầu và phỏng vấn stakeholder", level=2)
    add_body(doc, "Requirement elicitation là quá trình khám phá, làm rõ và kiểm chứng nhu cầu của các bên liên quan. Một cuộc phỏng vấn tốt không chỉ hỏi chức năng chính mà còn kiểm tra điều kiện, ngoại lệ, dữ liệu, quyền hạn và yêu cầu phi chức năng.")
    doc.add_heading("2.2. Mô phỏng stakeholder bằng LLM", level=2)
    add_body(doc, "LLM chịu trách nhiệm diễn đạt câu trả lời tự nhiên, trong khi policy và information gating giới hạn những facts mà stakeholder được phép tiết lộ. Cách tách này giúp giảm rò rỉ ground truth và tạo điều kiện kiểm thử.")
    doc.add_heading("2.3. Đánh giá requirement", level=2)
    add_body(doc, "Sau phiên chat, transcript được trích xuất thành requirement, chuẩn hóa và so khớp với hidden requirement của scenario. Coverage chỉ là chỉ số về mức độ khai thác requirement; nó không tự động chứng minh sinh viên giao tiếp tốt hơn.")

    # Chapter 3
    doc.add_heading("CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", level=1)
    doc.add_heading("3.1. Kiến trúc tổng thể", level=2)
    add_table(doc, ["Thành phần", "Công nghệ", "Trách nhiệm"], [["Frontend", "Vite + TypeScript", "Giao diện student, lecturer, admin"], ["Backend", "ASP.NET Core 9", "Auth, session, transaction, API"], ["AI service", "FastAPI + Gemini", "Persona, gating, extraction, evaluation"], ["Database", "PostgreSQL", "Người dùng, session, scenario, evaluation"]])
    add_caption(doc, "Hình 3.1. Kiến trúc ba dịch vụ của ReqSimulator (bản minh họa)")
    add_body(doc, "Backend là ranh giới công khai và nơi lưu dữ liệu. AI service được tách riêng để cô lập provider/model, logic prompt và các tác vụ xử lý ngôn ngữ.")
    doc.add_heading("3.2. Luồng phỏng vấn", level=2)
    add_table(doc, ["Bước", "Xử lý"], [["1", "Sinh viên chọn scenario và tạo session"], ["2", "Backend gửi câu hỏi cùng snapshot scenario cho AI service"], ["3", "Gating quyết định facts được phép sử dụng"], ["4", "LLM sinh câu trả lời và consistency guard hậu kiểm"], ["5", "Backend lưu transcript và trạng thái phiên"]])
    doc.add_heading("3.3. Các yêu cầu phi chức năng", level=2)
    add_bullet(doc, "Không để lộ hidden requirement ngoài policy của scenario.")
    add_bullet(doc, "Không lưu điểm 0 giả khi AI hoặc provider bị lỗi.")
    add_bullet(doc, "Session cũ phải giữ nguyên scenario version đã bắt đầu.")

    # Chapter 4
    doc.add_heading("CHƯƠNG 4. XÂY DỰNG VÀ ĐÁNH GIÁ HỆ THỐNG", level=1)
    doc.add_heading("4.1. Triển khai các module chính", level=2)
    add_body(doc, "Frontend cung cấp các màn hình đăng nhập, chọn scenario, chat, xem kết quả và quản trị. Backend xử lý xác thực JWT, phân quyền, session, versioning và lecturer review. AI service thực hiện persona response, gating, extraction và evaluation.")
    doc.add_heading("4.2. Pipeline đánh giá", level=2)
    add_table(doc, ["Giai đoạn", "Đầu vào", "Đầu ra"], [["Extraction", "Transcript", "Danh sách requirement"], ["Normalization", "Requirement thô", "Requirement chuẩn hóa"], ["Matching", "Requirement + ground truth", "Matched/partial/missing"], ["Coverage", "Kết quả matching", "Coverage score"], ["Feedback", "Coverage + missing", "Phản hồi học tập"]])
    doc.add_heading("4.3. Baseline kiểm thử", level=2)
    add_table(doc, ["Nhóm kiểm tra", "Trạng thái mẫu"], [["Backend Release build", "Thành công"], ["AI service unittest", "46/46 tại baseline"], ["Frontend test", "Còn giới hạn"], ["Integration runtime", "Cần PostgreSQL test biệt lập"]])
    add_body(doc, "Các con số trên chỉ là baseline kỹ thuật, không phải bằng chứng về hiệu quả học tập. Muốn kết luận về hiệu quả giáo dục cần pilot, rubric và dữ liệu được gắn nhãn.")

    # Chapter 5
    doc.add_heading("CHƯƠNG 5. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1)
    doc.add_heading("5.1. Kết quả đạt được", level=2)
    add_body(doc, "Bản thiết kế đã hình thành một luồng end-to-end gồm scenario versioning, stakeholder simulation, gating, transcript persistence và requirement-level evaluation.")
    doc.add_heading("5.2. Hạn chế", level=2)
    add_bullet(doc, "Matching một-một và normalization cần được kiểm chứng đầy đủ trên runtime.")
    add_bullet(doc, "Frontend và integration test chưa cân bằng với backend/AI test.")
    add_bullet(doc, "Hiệu quả giáo dục và mức độ giống người chưa được chứng minh bằng user study.")
    doc.add_heading("5.3. Hướng phát triển", level=2)
    add_body(doc, "Ưu tiên tiếp theo là hoàn thiện structured extraction, normalization, one-to-one matching, coverage và learning feedback; sau đó mới mở rộng dual-loop adaptive learning và nghiên cứu với sinh viên.")

    doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
    add_body(doc, "[1] Tài liệu tự nghiên cứu ReqSimulator, D:\\KLTN\\docs\\research_academic\\BAO_CAO_TU_NGHIEN_CUU_REQSIMULATOR.md.")
    add_body(doc, "[2] Bản đồ codebase ReqSimulator, D:\\KLTN\\docs\\project_map.md.")
    add_body(doc, "[3] Đặc tả Use Case và UML của ReqSimulator, D:\\KLTN\\docs\\architecture_specs\\REQSIMULATOR_UML_USE_CASE_SPEC.md.")

    doc.core_properties.title = "Mẫu khóa luận ReqSimulator"
    doc.core_properties.subject = "Bản mẫu bố cục đơn giản"
    doc.core_properties.author = "ReqSimulator Team"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
